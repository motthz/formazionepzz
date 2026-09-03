"""Crea 9 template di esempio (1 docx per reparto + 1 TUTTI + 1 xlsx)."""

from docx import Document
from openpyxl import Workbook
from pathlib import Path

ROOT = Path(__file__).resolve().parent
TEMPL_DIR = ROOT / "templates"
TEMPL_DIR.mkdir(exist_ok=True)


DOCX_TEMPLATES = [
    ("AMMINISTRAZIONE", 1, "AMM", "Richiesta materiale ufficio per *nome* — data *data*",
     """Gentile ufficio amministrazione,
si richiede la fornitura di cancelleria e materiale informatico standard per il/la sottoscritto/a *nome*,
inserito/a in data *data* presso la sede.
Cordiali saluti.
"""),
    ("COMMERCIALE", 2, "COM", "Scheda cliente primo contatto *nome* del *data*",
     """Dati anagrafici:
Nome dipendente: *nome*
Data ingresso: *data*
Reparto: COMMERCIALE

Note primo giorno:
- Presentazione al responsabile commerciale
- Assegnazione codice CRM
- Visione listino prezzi base
"""),
    ("LOGISTICA", 2, "LOG", "Checklist ricevimento merci per *nome* del *data*",
     """CHECKLIST OPERATIVO LOGISTICA
Operatore: *nome*
Data inizio: *data*

□ Conoscenza ubicazione banchine
□ Verifica DDT e fatture
□ Registro entrate/uscite
□ Procedure di carico/scarico con muletto
"""),
    ("MAGAZZINO", 2, "MAG", "Registro inventario iniziale *nome* *data*",
     """REGISTRO MAGAZZINO
Addetto: *nome*
Data presa servizio: *data*

- Inventario reparto 1 (materie prime)
- Inventario reparto 2 (semilavorati)
- Inventario reparto 3 (prodotti finiti)
- Ubicazione scaffalature e codici a barre
"""),
    ("PRODUZIONE", 3, "PRO", "Istruzioni di lavoro linea produzione *nome* *data*",
     """ISTRUZIONI LAVORO — LINEA PRODUZIONE
Operatore: *nome*
Data: *data*

1. Avvio linea e checklist giornaliera
2. Controllo qualità primo pezzo
3. Registrazione produzione oraria
4. Pulizia postazione e fine turno
"""),
    ("RISORSE UMANE", 1, "UMA", "Scheda personale *nome* — *data*",
     """SCHEDA PERSONALE — INSERIMENTO
Dipendente: *nome*
Data assunzione: *data*

□ Conferma dati anagrafici
□ Contratto sottoscritto
□ Informativa privacy GDPR
□ Consegna badge e chiavi
□ Inserimento software presenze
"""),
    ("SICUREZZA", 2, "SIC", "Scheda di formazione sicurezza generale — *nome* *data*",
     """FORMAZIONE GENERALE SULLA SICUREZZA LAVORATORI
Lavoratore: *nome*
Data corso: *data*

Argomenti trattati:
- Rischi generali di azienda
- DPI di reparto
- Uscite di emergenza e punto di raccolta
- Primo soccorso e antincendio
"""),
    ("VENDITE", 2, "VEN", "Scheda obiettivi vendite trimestrali *nome* del *data*",
     """PIANO VENDITE PERSONALIZZATO
Addetto vendite: *nome*
Data inizio: *data*

Obiettivo mensile: € ____________
Target nuovi clienti: ____________
Prodotti da promuovere: ____________
Cliente tutor assegnato: ____________
"""),
    ("TUTTI", 1, "GEN", "Regolamento generale aziendale per *nome* del *data*",
     """REGOLAMENTO GENERALE AZIENDALE
Dipendente: *nome*
Data presa visione: *data*

Il/la sottoscritto/a dichiara di aver ricevuto, letto e compreso il Regolamento Interno
dell'Azienda, le Norme sulla Sicurezza e sul comportamento nei luoghi di lavoro,
nonché la Privacy Policy vigente.
Firmato digitalmente in data odierna.
"""),
]


XLSX_TEMPLATES = [
    ("MAGAZZINO", 1, "MAG", "Registro carico scarico di magazzino"),
]


def make_docx(reparto: str, copie: int, codice: str, title: str, body: str):
    doc = Document()
    doc.add_heading(title, level=1)
    for para in body.split("\n"):
        doc.add_paragraph(para)
    table = doc.add_table(rows=3, cols=2)
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        pass
    table.rows[0].cells[0].text = "Dipendente"
    table.rows[0].cells[1].text = "*nome*"
    table.rows[1].cells[0].text = "Data di ingresso"
    table.rows[1].cells[1].text = "*data*"
    table.rows[2].cells[0].text = "Reparto"
    table.rows[2].cells[1].text = reparto.upper()
    fname = f"{reparto}_{copie}_{codice}.docx"
    doc.save(str(TEMPL_DIR / fname))
    return fname


def make_xlsx(reparto: str, copie: int, codice: str, title: str):
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]
    ws["A1"] = title.upper()
    ws.merge_cells("A1:D1")
    ws["A3"] = "Dipendente"
    ws["B3"] = "*nome*"
    ws["A4"] = "Data ingresso"
    ws["B4"] = "*data*"
    ws["A5"] = "Reparto"
    ws["B5"] = reparto.upper()
    ws["A7"] = "Movimento"
    ws["B7"] = "Codice articolo"
    ws["C7"] = "Quantità"
    ws["D7"] = "Firma operatore (*nome*)"
    fname = f"{reparto}_{copie}_{codice}.xlsx"
    wb.save(str(TEMPL_DIR / fname))
    return fname


creati = 0
for rep, copie, cod, titolo, corpo in DOCX_TEMPLATES:
    fn = make_docx(rep, copie, cod, titolo, corpo)
    print(f"  · {fn}")
    creati += 1

for rep, copie, cod, titolo in XLSX_TEMPLATES:
    fn = make_xlsx(rep, copie, cod, titolo)
    print(f"  · {fn}")
    creati += 1

print(f"\nCreati {creati} template di esempio in {TEMPL_DIR}")
