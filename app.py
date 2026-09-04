"""Formazioni PZZ: crea dossier PDF locali a partire da template Word ed Excel."""

from __future__ import annotations

import calendar
import copy as _copy_mod
import csv
import ctypes
import functools
import hashlib
import json
import os
import platform
import queue
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import traceback
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any

try:
    import tkinter as tk
    from tkinter import BOTH, END, LEFT, RIGHT, X, BooleanVar, StringVar, filedialog, messagebox, ttk
except ImportError:
    tk = None  # type: ignore[assignment]
    BOTH = END = LEFT = RIGHT = X = None  # type: ignore[assignment]
    BooleanVar = StringVar = filedialog = messagebox = ttk = None  # type: ignore[assignment]
from xml.sax.saxutils import escape

from docx import Document
from openpyxl import load_workbook
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    KeepInFrame,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


def _resolve_app_dir() -> Path:
    if getattr(sys, "frozen", False):
        exe_dir = Path(sys.executable).resolve().parent
        if exe_dir.name.lower() == "release":
            return exe_dir.parent
        return exe_dir
    return Path(__file__).resolve().parent


APP_DIR = _resolve_app_dir()
DEFAULT_TEMPLATE_DIR = APP_DIR / "templates"
DEFAULT_OUTPUT_DIR = APP_DIR / "output"
HISTORY_FILE = APP_DIR / ".formazioni_history.json"
DEPARTMENTS_FILE = APP_DIR / "reparti.txt"
SUPPORTED_EXTENSIONS = {".doc", ".docx", ".xls", ".xlsx", ".pdf"}
ALL_DEPARTMENT_NAMES = {"TUTTI", "TUTTE", "ALL"}
FILENAME_PATTERN = re.compile(
    r"^(?P<department>.+)_(?P<count>\d+)_(?P<code>[A-Za-z]{2,5})$",
    re.IGNORECASE,
)
SETTINGS_FILE = APP_DIR / "settings.json"
HASHES_FILE = APP_DIR / ".template_hashes.json"
LANG_DIR = APP_DIR / "lang"
DEFAULT_LANG = "it"
DEFAULT_THEME = "light"
MONTH_KEYS = (
    "dp_jan", "dp_feb", "dp_mar", "dp_apr", "dp_may", "dp_jun",
    "dp_jul", "dp_aug", "dp_sep", "dp_oct", "dp_nov", "dp_dec",
)

_STYLE_CACHE: dict[str, object] | None = None
_TEMPLATE_STORY_CACHE: dict[tuple, list[object]] = {}


# --------------------------- SETTINGS -------------------------------------

def _default_settings() -> dict[str, Any]:
    return {
        "theme": DEFAULT_THEME,
        "language": DEFAULT_LANG,
        "last_template_dir": str(DEFAULT_TEMPLATE_DIR),
        "last_output_dir": str(DEFAULT_OUTPUT_DIR),
    }


def load_settings() -> dict[str, Any]:
    data = _default_settings()
    if SETTINGS_FILE.exists():
        try:
            raw = json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
            if isinstance(raw, dict):
                data.update({k: v for k, v in raw.items() if k in data})
        except (OSError, json.JSONDecodeError):
            pass
    return data


def save_settings(data: dict[str, Any]) -> None:
    try:
        APP_DIR.mkdir(parents=True, exist_ok=True)
        SETTINGS_FILE.write_text(
            json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except OSError:
        pass


# --------------------------- I18N -----------------------------------------

def available_languages() -> list[str]:
    if not LANG_DIR.exists():
        return [DEFAULT_LANG]
    return sorted(p.stem for p in LANG_DIR.glob("*.json"))


def load_language(code: str) -> dict[str, str]:
    fallback: dict[str, str] = {}
    if (LANG_DIR / f"{DEFAULT_LANG}.json").exists():
        try:
            fallback = json.loads(
                (LANG_DIR / f"{DEFAULT_LANG}.json").read_text(encoding="utf-8")
            )
        except (OSError, json.JSONDecodeError):
            fallback = {}
    chosen_file = LANG_DIR / f"{code}.json"
    chosen: dict[str, str] = {}
    if chosen_file.exists():
        try:
            chosen = json.loads(chosen_file.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            chosen = {}
    merged = dict(fallback)
    merged.update(chosen)
    return merged


# --------------------------- HASHING --------------------------------------

def compute_template_hash(path: Path) -> str:
    hasher = hashlib.md5()
    try:
        with path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(65536), b""):
                hasher.update(chunk)
    except OSError:
        return ""
    return hasher.hexdigest()


def load_saved_hashes() -> dict[str, str]:
    if not HASHES_FILE.exists():
        return {}
    try:
        data = json.loads(HASHES_FILE.read_text(encoding="utf-8"))
        return {str(k): str(v) for k, v in data.items()} if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {}


def save_hashes(data: dict[str, str]) -> None:
    try:
        HASHES_FILE.write_text(
            json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except OSError:
        pass


def classify_template_hashes(
    templates: list[Any], saved: dict[str, str]
) -> dict[str, str]:
    result: dict[str, str] = {}
    for template in templates:
        key = str(template.path)
        current = compute_template_hash(template.path)
        previous = saved.get(key)
        if previous is None:
            result[key] = "new"
        elif previous == current:
            result[key] = "ok"
        else:
            result[key] = "modified"
    return result


def clear_caches() -> None:
    global _STYLE_CACHE
    _STYLE_CACHE = None
    _TEMPLATE_STORY_CACHE.clear()
    _office_command.cache_clear()


# --------------------------- BUSINESS -------------------------------------

def load_departments_from_file() -> list[str]:
    if not DEPARTMENTS_FILE.exists():
        return []
    try:
        raw = DEPARTMENTS_FILE.read_text(encoding="utf-8")
    except OSError:
        return []
    departments: list[str] = []
    seen: set[str] = set()
    for line in raw.splitlines():
        name = line.strip()
        if not name or name.startswith("#"):
            continue
        key = name.upper()
        if key not in seen:
            seen.add(key)
            departments.append(key)
    return departments


@dataclass(frozen=True)
class TemplateFile:
    path: Path
    department: str
    copies: int
    code: str

    @property
    def display_name(self) -> str:
        return self.path.name

    @property
    def is_for_every_department(self) -> bool:
        return self.department.upper() in ALL_DEPARTMENT_NAMES


def parse_template(path: Path) -> TemplateFile | None:
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        return None
    match = FILENAME_PATTERN.fullmatch(path.stem)
    if not match:
        return None
    return TemplateFile(
        path=path,
        department=match.group("department"),
        copies=int(match.group("count")),
        code=match.group("code").upper(),
    )


def discover_templates(folder: Path) -> tuple[list[TemplateFile], list[Path]]:
    valid: list[TemplateFile] = []
    ignored: list[Path] = []
    if not folder.exists():
        return valid, ignored
    for path in sorted(folder.rglob("*")):
        if path.is_file() and path.name.lower() != "readme.md":
            template = parse_template(path)
            if template and template.copies > 0:
                valid.append(template)
            elif path.suffix.lower() in SUPPORTED_EXTENSIONS:
                ignored.append(path)
    return valid, ignored


def department_options(templates: list[TemplateFile]) -> list[str]:
    from_file = load_departments_from_file()
    from_templates = {
        template.department.upper()
        for template in templates
        if not template.is_for_every_department
    }
    merged: list[str] = []
    seen: set[str] = set()
    for name in from_file:
        if name not in seen:
            seen.add(name)
            merged.append(name)
    for name in sorted(from_templates):
        if name not in seen:
            seen.add(name)
            merged.append(name)
    return merged


def templates_for_department(
    templates: list[TemplateFile], department: str
) -> list[TemplateFile]:
    chosen = department.strip().upper()
    matching = [
        template
        for template in templates
        if template.is_for_every_department
        or template.department.upper() == chosen
    ]
    return sorted(
        matching,
        key=lambda item: (
            item.is_for_every_department is False,
            item.department.upper(),
            item.path.name.lower(),
        ),
    )


def templates_for_departments(
    templates: list[TemplateFile], departments: list[str]
) -> list[TemplateFile]:
    chosen = {d.strip().upper() for d in departments if d and d.strip()}
    seen_paths: set[Path] = set()
    result: list[TemplateFile] = []
    for template in templates:
        key = template.path
        if key in seen_paths:
            continue
        if template.is_for_every_department or (template.department.upper() in chosen):
            seen_paths.add(key)
            result.append(template)
    return sorted(
        result,
        key=lambda item: (
            item.is_for_every_department is False,
            item.department.upper(),
            item.path.name.lower(),
        ),
    )


def replace_placeholders(value: object, employee_name: str, entry_date: str) -> object:
    if not isinstance(value, str):
        return value
    value = re.sub(r"\*nome\*", employee_name, value, flags=re.IGNORECASE)
    return re.sub(r"\*data\*", entry_date, value, flags=re.IGNORECASE)


def replace_paragraph(paragraph, employee_name: str, entry_date: str) -> None:
    original = paragraph.text
    for run in paragraph.runs:
        run.text = str(replace_placeholders(run.text, employee_name, entry_date))
    if paragraph.text != original:
        return
    replaced = replace_placeholders(original, employee_name, entry_date)
    if original == replaced:
        return
    if paragraph.runs:
        paragraph.runs[0].text = str(replaced)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(str(replaced))


def replace_docx_placeholders(document: Document, employee_name: str, entry_date: str) -> None:
    for paragraph in document.paragraphs:
        replace_paragraph(paragraph, employee_name, entry_date)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph(paragraph, employee_name, entry_date)
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                replace_paragraph(paragraph, employee_name, entry_date)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_paragraph(paragraph, employee_name, entry_date)


def replace_xlsx_placeholders(workbook, employee_name: str, entry_date: str) -> None:
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                try:
                    cell.value = replace_placeholders(cell.value, employee_name, entry_date)
                except AttributeError:
                    pass


@functools.lru_cache(maxsize=1)
def _office_command() -> str | None:
    for command in ("libreoffice", "soffice"):
        if shutil.which(command):
            return command
    if os.name == "nt":
        try:
            import win32com.client  # type: ignore[import-not-found]  # noqa: F401
        except ImportError:
            return None
        return "microsoft-office"
    return None


def _convert_with_office(source: Path, output_dir: Path, extension: str) -> Path:
    command = _office_command()
    if command is None:
        raise RuntimeError(
            "Per mantenere identico il layout dei template Word/Excel serve "
            "LibreOffice installato e disponibile nel PATH."
        )
    output_dir.mkdir(parents=True, exist_ok=True)
    converted = output_dir / f"{source.stem}.{extension}"
    if command == "microsoft-office":
        import win32com.client  # type: ignore[import-not-found]

        if source.suffix.lower() in {".doc", ".docx"}:
            word = win32com.client.DispatchEx("Word.Application")
            document = word.Documents.Open(str(source.resolve()))
            try:
                if extension == "pdf":
                    document.ExportAsFixedFormat(str(converted.resolve()), 17)
                else:
                    document.SaveAs2(str(converted.resolve()), FileFormat=16)
            finally:
                document.Close(False)
                word.Quit()
        else:
            excel = win32com.client.DispatchEx("Excel.Application")
            workbook = excel.Workbooks.Open(str(source.resolve()))
            try:
                if extension == "pdf":
                    workbook.ExportAsFixedFormat(0, str(converted.resolve()))
                else:
                    workbook.SaveAs(str(converted.resolve()), FileFormat=51)
            finally:
                workbook.Close(False)
                excel.Quit()
        if not converted.exists():
            raise RuntimeError(f"Conversione Microsoft Office fallita per {source.name}")
        return converted

    profile = output_dir / "office-profile"
    profile_uri = profile.resolve().as_uri()
    result = subprocess.run(
        [
            command,
            "--headless",
            "--convert-to",
            extension,
            "--outdir",
            str(output_dir),
            f"-env:UserInstallation={profile_uri}",
            str(source),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0 or not converted.exists():
        detail = (result.stderr or result.stdout).strip()
        raise RuntimeError(f"Conversione Office fallita per {source.name}: {detail}")
    return converted


def _convert_with_office_batch(
    sources: list[Path], output_dir: Path, extension: str
) -> dict[Path, Path]:
    if not sources:
        return {}
    command = _office_command()
    if command is None:
        raise RuntimeError("Motore Office non disponibile per la conversione dei template.")
    output_dir.mkdir(parents=True, exist_ok=True)
    if command == "microsoft-office":
        import win32com.client  # type: ignore[import-not-found]

        applications = {}
        converted: dict[Path, Path] = {}
        try:
            for source in sources:
                kind = "word" if source.suffix.lower() == ".docx" else "excel"
                if kind not in applications:
                    applications[kind] = (
                        win32com.client.DispatchEx("Word.Application")
                        if kind == "word"
                        else win32com.client.DispatchEx("Excel.Application")
                    )
                application = applications[kind]
                target = output_dir / f"{source.stem}.{extension}"
                document = (
                    application.Documents.Open(str(source.resolve()))
                    if kind == "word"
                    else application.Workbooks.Open(str(source.resolve()))
                )
                try:
                    if extension == "pdf":
                        if kind == "word":
                            document.ExportAsFixedFormat(str(target.resolve()), 17)
                        else:
                            document.ExportAsFixedFormat(0, str(target.resolve()))
                finally:
                    document.Close(False)
                converted[source] = target
        finally:
            for application in applications.values():
                application.Quit()
        return converted

    profile = output_dir / "office-profile"
    result = subprocess.run(
        [
            command,
            "--headless",
            "--convert-to",
            extension,
            "--outdir",
            str(output_dir),
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            *[str(source) for source in sources],
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    converted = {source: output_dir / f"{source.stem}.{extension}" for source in sources}
    if result.returncode != 0 or any(not path.exists() for path in converted.values()):
        detail = (result.stderr or result.stdout).strip()
        raise RuntimeError(f"Conversione Office fallita: {detail}")
    return converted


def _prepare_office_template(
    template: Path,
    work_dir: Path,
    employee_name: str,
    entry_date: str,
) -> Path:
    suffix = template.suffix.lower()
    work_dir.mkdir(parents=True, exist_ok=True)
    if suffix in {".doc", ".xls"}:
        modern_extension = "docx" if suffix == ".doc" else "xlsx"
        source = _convert_with_office(template, work_dir, modern_extension)
    else:
        source = work_dir / template.name
        shutil.copy2(template, source)

    if source.suffix.lower() == ".docx":
        document = Document(str(source))
        replace_docx_placeholders(document, employee_name, entry_date)
        document.save(str(source))
    elif source.suffix.lower() == ".xlsx":
        workbook = load_workbook(source, data_only=False, read_only=False)
        try:
            replace_xlsx_placeholders(workbook, employee_name, entry_date)
            workbook.save(source)
        finally:
            workbook.close()
    return source


def _merge_pdfs(output_path: Path, pdf_paths: list[Path]) -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    for pdf_path in pdf_paths:
        writer.append(str(pdf_path))
    with output_path.open("wb") as handle:
        writer.write(handle)


def _remove_trailing_blank_pages(pdf_path: Path) -> Path:
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(str(pdf_path))
    last_page = len(reader.pages)
    while last_page > 1:
        page = reader.pages[last_page - 1]
        text = page.extract_text() or ""
        resources = page.get("/Resources")
        if resources is not None:
            resources = resources.get_object()
        has_images = bool(resources and resources.get("/XObject"))
        content = page.get_contents()
        has_drawing_commands = bool(content and content.get_data().strip())
        if text.strip() or has_images or has_drawing_commands:
            break
        last_page -= 1
    pages = reader.pages[:last_page]
    if len(pages) == 1:
        return pdf_path

    trimmed = pdf_path.with_name(f"{pdf_path.stem}-trimmed.pdf")
    writer = PdfWriter()
    for page in pages:
        writer.add_page(page)
    with trimmed.open("wb") as handle:
        writer.write(handle)
    return trimmed


def _pdf_is_landscape(pdf_path: Path) -> bool:
    from pypdf import PdfReader

    reader = PdfReader(str(pdf_path))
    return bool(reader.pages and reader.pages[0].mediabox.width > reader.pages[0].mediabox.height)


def _build_pdf_native(
    output_path: Path,
    employee_name: str,
    entry_date: str,
    department: str,
    role: str,
    notes: str,
    expanded: list[tuple[TemplateFile, int]],
    progress_cb=None,
) -> int:
    total_steps = max(1, len({t.path for t, _ in expanded}) + 1)
    step = 0
    with tempfile.TemporaryDirectory(prefix="formazioni-pdf-") as temp_name:
        temp_dir = Path(temp_name)
        prepared: dict[Path, Path] = {}

        unique_templates = {}
        for index, (template, _) in enumerate(expanded):
            if template.path.suffix.lower() != ".pdf" and template.path not in prepared:
                unique_templates[template.path] = index

        for template_path, index in unique_templates.items():
            prepared[template_path] = _prepare_office_template(
                template_path, temp_dir / f"template-{index}", employee_name, entry_date
            )
            step += 1
            if progress_cb:
                progress_cb(step, total_steps)

        if prepared:
            converted = _convert_with_office_batch(
                list(prepared.values()), temp_dir / "converted", "pdf"
            )
        else:
            converted = {}
        step += 1
        if progress_cb:
            progress_cb(step, total_steps)

        pages = [
            _remove_trailing_blank_pages(template.path)
            if template.path.suffix.lower() == ".pdf"
            else _remove_trailing_blank_pages(converted[prepared[template.path]])
            for template, _copy_number in expanded
        ]
        _merge_pdfs(output_path, pages)
    return len(expanded)


def paragraph_text(text: str, style: ParagraphStyle) -> Paragraph:
    safe = escape(text).replace("\n", "<br/>")
    return Paragraph(safe or " ", style)


TABLE_WIDTH = 174 * mm


def _table_widths(values: list[list[str]], available_width: float = TABLE_WIDTH) -> list[float]:
    column_count = max((len(row) for row in values), default=0)
    if not column_count:
        return []
    lengths = []
    for column in range(column_count):
        longest = max(
            (len(str(row[column]).replace("\n", " ").strip()) for row in values if column < len(row)),
            default=4,
        )
        lengths.append(max(4, min(longest, 36)))
    minimum = available_width / column_count
    weights = [max(length, 8) for length in lengths]
    total = sum(weights)
    widths = [available_width * weight / total for weight in weights]
    if any(width < minimum for width in widths):
        widths = [minimum] * column_count
    return widths


def _table_flowable(rows: list[list[object]], raw_values: list[list[str]]) -> Table:
    column_count = max((len(row) for row in rows), default=0)
    rows = [row + [""] * (column_count - len(row)) for row in rows]
    raw_values = [row + [""] * (column_count - len(row)) for row in raw_values]
    return Table(
        rows,
        colWidths=_table_widths(raw_values),
        repeatRows=1,
        splitByRow=1,
        splitInRow=0,
        hAlign="LEFT",
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7eef1")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#173642")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#b9c9cf")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        ),
    )


def docx_story(
    path: Path,
    employee_name: str,
    entry_date: str,
    styles: dict[str, ParagraphStyle],
) -> list[object]:
    document = Document(path)
    replace_docx_placeholders(document, employee_name, entry_date)
    story: list[object] = []
    style_name_cache: dict[int, str] = {}
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        heading = ""
        try:
            if paragraph.style is not None:
                pid = id(paragraph.style)
                if pid not in style_name_cache:
                    style_name_cache[pid] = (paragraph.style.name or "").lower()
                heading = style_name_cache[pid]
        except Exception:
            heading = ""
        style = styles["subheading"] if "heading" in heading else styles["body"]
        story.append(paragraph_text(text, style))
        story.append(Spacer(1, 2.4 * mm))
    for table in document.tables:
        rows: list[list[object]] = []
        raw_rows: list[list[str]] = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                while values and not values[-1]:
                    values.pop()
                raw_rows.append(values)
                rows.append([paragraph_text(value, styles["table"]) for value in values])
        if rows:
            story.append(Spacer(1, 2 * mm))
            story.append(_table_flowable(rows, raw_rows))
            story.append(Spacer(1, 4 * mm))
    return story or [paragraph_text("Documento senza contenuto testuale.", styles["muted"])]


def xlsx_story(
    path: Path,
    employee_name: str,
    entry_date: str,
    styles: dict[str, ParagraphStyle],
) -> list[object]:
    workbook = load_workbook(path, data_only=False, read_only=False)
    story: list[object] = []
    try:
        for sheet in workbook.worksheets:
            rows: list[list[object]] = []
            raw_rows: list[list[str]] = []
            for row in sheet.iter_rows(values_only=True):
                values = [
                    replace_placeholders(value, employee_name, entry_date)
                    for value in row
                ]
                if any(value not in (None, "") for value in values):
                    values = [str(value) if value is not None else "" for value in values]
                    while values and not values[-1]:
                        values.pop()
                    raw_rows.append(values)
                    rows.append(
                        [paragraph_text(value, styles["table"]) for value in values]
                    )
            if rows:
                story.append(Paragraph(escape(sheet.title), styles["subheading"]))
                story.append(Spacer(1, 2 * mm))
                story.append(_table_flowable(rows, raw_rows))
                story.append(Spacer(1, 5 * mm))
    finally:
        workbook.close()
    return story or [paragraph_text("Foglio senza contenuto.", styles["muted"])]


def make_styles() -> dict[str, ParagraphStyle]:
    global _STYLE_CACHE
    if _STYLE_CACHE is not None:
        return _STYLE_CACHE  # type: ignore[return-value]
    base = getSampleStyleSheet()
    result = {
        "cover_title": ParagraphStyle("CoverTitle", parent=base["Title"], fontName="Helvetica-Bold",
                                     fontSize=25, leading=30, textColor=colors.HexColor("#173642"),
                                     alignment=TA_CENTER, spaceAfter=8 * mm),
        "cover_subtitle": ParagraphStyle("CoverSubtitle", parent=base["Normal"], fontName="Helvetica",
                                         fontSize=12, leading=17, textColor=colors.HexColor("#48636d"),
                                         alignment=TA_CENTER),
        "meta": ParagraphStyle("Meta", parent=base["Normal"], fontName="Helvetica",
                               fontSize=10.5, leading=16, textColor=colors.HexColor("#173642")),
        "heading": ParagraphStyle("Heading", parent=base["Heading1"], fontName="Helvetica-Bold",
                                  fontSize=15, leading=19, textColor=colors.HexColor("#173642"),
                                  spaceAfter=4 * mm),
        "subheading": ParagraphStyle("Subheading", parent=base["Heading2"], fontName="Helvetica-Bold",
                                     fontSize=11.5, leading=15, textColor=colors.HexColor("#247b7b"),
                                     spaceBefore=2 * mm, spaceAfter=2 * mm),
        "body": ParagraphStyle("Body", parent=base["BodyText"], fontName="Helvetica",
                               fontSize=10, leading=14, textColor=colors.HexColor("#263f47")),
        "table": ParagraphStyle("Table", parent=base["BodyText"], fontName="Helvetica",
                                fontSize=8.2, leading=10, textColor=colors.HexColor("#263f47")),
        "muted": ParagraphStyle("Muted", parent=base["BodyText"], fontName="Helvetica-Oblique",
                                fontSize=9.5, leading=13, textColor=colors.HexColor("#71838a")),
        "small": ParagraphStyle("Small", parent=base["BodyText"], fontName="Helvetica",
                                fontSize=8.5, leading=11, textColor=colors.HexColor("#71838a")),
        "pdf_page": ParagraphStyle("PdfPage", parent=base["BodyText"], fontName="Helvetica",
                                   fontSize=8.5, leading=11.5, textColor=colors.HexColor("#263f47")),
    }
    _STYLE_CACHE = result  # type: ignore[assignment]
    return result


def pdf_story(
    path: Path,
    employee_name: str,
    entry_date: str,
    styles: dict[str, ParagraphStyle],
    frame_width: float = 174 * mm,
    frame_height: float = 263 * mm,
) -> list[object]:
    story: list[object] = []
    pdf_style = styles["pdf_page"]
    is_landscape = frame_width > frame_height
    if is_landscape:
        pdf_style = ParagraphStyle("PdfPageLandscape", parent=styles["pdf_page"],
                                   fontSize=7.5, leading=10)
    try:
        from pypdf import PdfReader
    except Exception:
        story.append(paragraph_text(
            "Libreria pypdf non disponibile per leggere il template PDF.", styles["muted"]))
        return story
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        story.append(paragraph_text(f"Impossibile aprire il PDF: {exc}", styles["muted"]))
        return story
    if not reader.pages:
        story.append(paragraph_text("PDF senza pagine.", styles["muted"]))
        return story
    for idx, page in enumerate(reader.pages, start=1):
        page_story: list[object] = []
        if len(reader.pages) > 1:
            page_story.append(Paragraph(f"Pagina {idx} del PDF", styles["subheading"]))
        raw = ""
        try:
            raw = page.extract_text() or ""
        except Exception:
            raw = ""
        raw = replace_placeholders(raw, employee_name, entry_date)
        lines = [ln.rstrip() for ln in raw.splitlines()]
        if not any(ln.strip() for ln in lines):
            page_story.append(paragraph_text(
                "(Nessun testo estraibile dal PDF — file incluso come riferimento.)",
                styles["muted"]))
            page_story.append(Spacer(1, 3 * mm))
        else:
            for line in lines:
                if not line.strip():
                    page_story.append(Spacer(1, 1.5 * mm))
                    continue
                page_story.append(Paragraph(escape(line), pdf_style))
            page_story.append(Spacer(1, 4 * mm))
        story.append(KeepInFrame(frame_width, frame_height, page_story, mode="shrink"))
        if idx < len(reader.pages):
            story.append(PageBreak())
    return story or [paragraph_text("Documento PDF senza contenuto testuale.", styles["muted"])]


def build_pdf(
    output_path: Path,
    employee_name: str,
    entry_date: str,
    department: str,
    role: str,
    notes: str,
    templates: list[TemplateFile],
    progress_cb=None,
) -> int:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = make_styles()
    expanded = [
        (template, copy_number)
        for template in templates
        for copy_number in range(1, template.copies + 1)
    ]
    needs_legacy_office = any(
        template.path.suffix.lower() in {".doc", ".xls"}
        for template, _ in expanded
    )
    if _office_command() and needs_legacy_office:
        return _build_pdf_native(
            output_path, employee_name, entry_date, department, role, notes, expanded,
            progress_cb=progress_cb,
        )
    has_landscape_pdf = any(
        template.path.suffix.lower() == ".pdf" and _pdf_is_landscape(template.path)
        for template, _ in expanded
    )
    page_size = landscape(A4) if has_landscape_pdf else A4
    page_width, page_height = page_size
    horizontal_frame = min(page_width - 28 * mm, 210 * mm)
    vertical_frame = min(page_height - 28 * mm, 139 * mm)
    h_margin = 10 * mm if has_landscape_pdf else 14 * mm
    v_margin = 10 * mm if has_landscape_pdf else 14 * mm

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=page_size,
        rightMargin=h_margin, leftMargin=h_margin, topMargin=v_margin, bottomMargin=v_margin,
        title=f"Dossier formazione - {employee_name}",
        author="Formazioni PZZ",
    )
    story: list[object] = []
    story_cache_local: dict[Path, list[object]] = {}
    total_steps = max(1, len(expanded) + 1)

    def _fetch_story(template: TemplateFile, step_idx: int) -> list[object]:
        if progress_cb:
            progress_cb(step_idx, total_steps)
        if template.path in story_cache_local:
            return story_cache_local[template.path]
        suffix = template.path.suffix.lower()
        try:
            mtime = template.path.stat().st_mtime_ns
        except OSError:
            mtime = 0
        key = (
            hash(template.path.resolve().as_posix()),
            mtime, employee_name, entry_date, suffix,
            int(horizontal_frame), int(vertical_frame),
        )
        cached_global = _TEMPLATE_STORY_CACHE.get(key)
        if cached_global is not None:
            story_cache_local[template.path] = cached_global
            return cached_global
        if suffix == ".docx":
            generated = docx_story(template.path, employee_name, entry_date, styles)
        elif suffix == ".xlsx":
            generated = xlsx_story(template.path, employee_name, entry_date, styles)
        elif suffix == ".pdf":
            generated = pdf_story(template.path, employee_name, entry_date, styles,
                                  horizontal_frame, vertical_frame)
        else:
            generated = []
        story_cache_local[template.path] = generated
        _TEMPLATE_STORY_CACHE[key] = generated
        return generated

    for index, (template, copy_number) in enumerate(expanded):
        if index > 0:
            story.append(PageBreak())
        cached = _fetch_story(template, index + 1)
        if len(cached) > 3 or isinstance(cached[0] if cached else None, Table):
            story.extend(_copy_mod.copy(cached))
        else:
            story.extend(cached)
    doc.build(story)
    if progress_cb:
        progress_cb(total_steps, total_steps)
    return len(expanded)


def safe_file_part(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9À-ÿ_-]+", "-", value.strip())
    return cleaned.strip("-_") or "persona"


def open_folder(path: Path) -> None:
    try:
        if platform.system() == "Windows":
            os.startfile(str(path))  # type: ignore[attr-defined]
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", str(path)])
        else:
            subprocess.Popen(["xdg-open", str(path)])
    except OSError:
        pass


# --------------------------- UI HELPERS -----------------------------------

class Tooltip:
    """Tooltip minimale con delay e auto-dismiss."""

    def __init__(self, widget: tk.Misc, text_getter, delay_ms: int = 600, auto_dismiss_ms: int = 8000):
        self.widget = widget
        self.text_getter = text_getter if callable(text_getter) else lambda: text_getter
        self.delay_ms = delay_ms
        self.auto_dismiss_ms = auto_dismiss_ms
        self._after_id: str | None = None
        self._dismiss_after_id: str | None = None
        self.tip: tk.Toplevel | None = None
        widget.bind("<Enter>", self._on_enter, add="+")
        widget.bind("<Leave>", self._on_leave, add="+")
        widget.bind("<Motion>", self._on_leave, add="+")
        widget.bind("<ButtonPress>", self._on_leave, add="+")
        widget.bind("<Destroy>", self._on_leave, add="+")

    def _on_enter(self, _evt=None):
        self._cancel()
        self._after_id = self.widget.after(self.delay_ms, self._show)

    def _on_leave(self, _evt=None):
        self._cancel()
        self._hide()

    def _cancel(self):
        if self._after_id is not None:
            try:
                self.widget.after_cancel(self._after_id)
            except Exception:
                pass
            self._after_id = None
        if self._dismiss_after_id is not None:
            try:
                self.widget.after_cancel(self._dismiss_after_id)
            except Exception:
                pass
            self._dismiss_after_id = None

    def _show(self):
        self._after_id = None
        text = str(self.text_getter() or "").strip()
        if not text:
            return
        self.tip = t = tk.Toplevel(self.widget)
        t.wm_overrideredirect(True)
        try:
            t.wm_attributes("-topmost", True)
        except Exception:
            pass
        try:
            t.configure(bg="#fff2a8")
        except Exception:
            pass
        frame = tk.Frame(t, bg="#fff2a8", highlightthickness=1,
                         highlightbackground="#b39800", padx=8, pady=4)
        frame.pack(fill=BOTH, expand=True)
        label = tk.Label(frame, text=text, justify="left",
                         bg="#fff2a8", fg="#2a1a00",
                         font=("Segoe UI", 9), wraplength=360)
        label.pack()
        self.widget.update_idletasks()
        try:
            x = self.widget.winfo_rootx() + 12
            y = self.widget.winfo_rooty() + self.widget.winfo_height() + 4
            t.wm_geometry(f"+{x}+{y}")
        except Exception:
            pass
        if self.auto_dismiss_ms > 0:
            self._dismiss_after_id = self.widget.after(self.auto_dismiss_ms, self._hide)

    def _hide(self):
        if self.tip is not None:
            try:
                self.tip.destroy()
            except Exception:
                pass
            self.tip = None


class DatePickerFrame(tk.Frame):
    """Date picker puro Tk: 3 Combobox (giorno/mese/anno) con validazione."""

    def __init__(self, master, language: dict[str, str], initial: date | None = None, **kwargs):
        super().__init__(master, **kwargs)
        self.language = language
        self._build()
        if initial is None:
            initial = date.today()
        self.set_date(initial)
        self._bind_change()

    def _build(self):
        self.columnconfigure(0, weight=0)
        self.columnconfigure(1, weight=0)
        self.columnconfigure(2, weight=0)
        today = date.today()
        year_values = [str(y) for y in range(today.year - 5, today.year + 6)]
        self.day_cb = ttk.Combobox(self, values=[str(d) for d in range(1, 32)],
                                   width=4, state="readonly")
        self.month_cb = ttk.Combobox(self, values=[self.language.get(k, k) for k in MONTH_KEYS],
                                     width=12, state="readonly")
        self.year_cb = ttk.Combobox(self, values=year_values, width=6, state="readonly")
        self.day_cb.grid(row=0, column=0, padx=(0, 6), sticky="w")
        self.month_cb.grid(row=0, column=1, padx=(0, 6), sticky="w")
        self.year_cb.grid(row=0, column=2, sticky="w")

    def _bind_change(self):
        def sync(_e=None):
            days_max = self._compute_days()
            current = self.day_cb.current() + 1
            day_vals = [str(d) for d in range(1, days_max + 1)]
            self.day_cb.configure(values=day_vals)
            if current > days_max:
                self.day_cb.current(days_max - 1)
            self._on_change()

        self.day_cb.bind("<<ComboboxSelected>>", sync)
        self.month_cb.bind("<<ComboboxSelected>>", sync)
        self.year_cb.bind("<<ComboboxSelected>>", sync)

    def _compute_days(self) -> int:
        m = self.month_cb.current() + 1
        try:
            y = int(self.year_cb.get()) if self.year_cb.get() else date.today().year
        except ValueError:
            y = date.today().year
        if m < 1 or m > 12:
            return 31
        return calendar.monthrange(y, m)[1]

    def _on_change(self):
        pass

    def get_date(self) -> date | None:
        try:
            d = int(self.day_cb.get())
            m = self.month_cb.current() + 1
            y = int(self.year_cb.get())
            if not (1 <= m <= 12):
                return None
            dm = calendar.monthrange(y, m)[1]
            if not (1 <= d <= dm):
                d = dm
            return date(y, m, d)
        except (ValueError, tk.TclError):
            return None

    def get_string(self) -> str:
        d = self.get_date()
        return d.strftime("%d/%m/%Y") if d else ""

    def set_date(self, d: date) -> None:
        try:
            self.day_cb.current(d.day - 1)
        except Exception:
            pass
        try:
            self.month_cb.current(d.month - 1)
        except Exception:
            pass
        vals = list(self.year_cb["values"])
        target = str(d.year)
        if target in vals:
            self.year_cb.current(vals.index(target))
        else:
            self.year_cb.set(target)

    def configure_language(self, language: dict[str, str]) -> None:
        self.language = language
        current_month_index = self.month_cb.current()
        self.month_cb.configure(values=[language.get(k, k) for k in MONTH_KEYS])
        if 0 <= current_month_index < 12:
            self.month_cb.current(current_month_index)


# --------------------------- APP ------------------------------------------

class FormazioniApp:
    def __init__(self, root: tk.Tk) -> None:
        if tk is None or ttk is None:
            raise RuntimeError(
                "Questa installazione di Python non include Tkinter. "
                "Installa Python con il supporto Tk e riavvia Formazioni PZZ."
            )
        self.root = root

        self.settings = load_settings()
        self.language_code = self.settings.get("language", DEFAULT_LANG)
        self.language = load_language(self.language_code)
        self.saved_hashes = load_saved_hashes()
        self.hash_status: dict[str, str] = {}

        self.template_dir = StringVar(value=self.settings.get("last_template_dir", str(DEFAULT_TEMPLATE_DIR)))
        self.output_dir = StringVar(value=self.settings.get("last_output_dir", str(DEFAULT_OUTPUT_DIR)))
        self.employee_name = StringVar()
        self.role = StringVar()
        self.department = StringVar()
        self.auto_open = BooleanVar(value=True)
        self.multi_dept_mode = BooleanVar(value=False)
        self.theme = StringVar(value=self.settings.get("theme", DEFAULT_THEME))
        self.language_var = StringVar(value=self.language_code)
        self.status = StringVar(value=self.tr("status_initial"))
        self.count_label = StringVar(value=self.tr("count_none"))
        self.progress_label = StringVar(value=self.tr("progress_ready"))
        self.hash_stat_label = StringVar(value="")
        self.templates: list[TemplateFile] = []
        self.ignored: list[Path] = []
        self.template_inclusion: dict[Path, bool] = {}
        self.multi_dept_values: dict[str, BooleanVar] = {}
        self._tooltips: list[Tooltip] = []
        self._queue: queue.Queue[tuple[str, Any]] = queue.Queue()
        self._worker_active = False
        self._current_width = 1120
        self._current_breakpoint = "wide"
        self._progressbar: ttk.Progressbar | None = None

        self.root.title(self.tr("app_title"))
        self.root.geometry("1120x820")
        self.root.minsize(980, 720)

        self._configure_style()
        self._apply_theme_root()
        self._build_scaffold()
        self._build_header()
        self._build_body()
        self._install_drain_loop()
        self.root.bind("<Configure>", self._on_root_resize)
        self.refresh_templates()

    # ----------------------------- Utilities ------------------------------
    def tr(self, key: str, **kwargs: Any) -> str:
        raw = self.language.get(key, key)
        try:
            return raw.format(**kwargs) if kwargs else raw
        except (KeyError, IndexError):
            return raw

    def _persist_settings(self) -> None:
        self.settings["theme"] = self.theme.get()
        self.settings["language"] = self.language_code
        self.settings["last_template_dir"] = self.template_dir.get()
        self.settings["last_output_dir"] = self.output_dir.get()
        save_settings(self.settings)

    # ----------------------------- Style ----------------------------------
    def _apply_theme_root(self):
        theme = self.theme.get()
        if theme == "dark":
            self.root.configure(bg="#0e1a22")
        else:
            self.root.configure(bg="#eef3f3")

    def _configure_style(self):
        dark = self.theme.get() == "dark"
        style = ttk.Style()
        try:
            style.theme_use("clam")
        except Exception:
            pass

        if dark:
            app_bg = "#0e1a22"
            card_bg = "#17242c"
            card_body_bg = "#17242c"
            sh1 = "#060d11"
            sh2 = "#0a171d"
            text = "#eaf0f2"
            text_muted = "#95adb6"
            gold = "#d9a13f"
            accent_bg = "#07131b"
            accent_fg = "#eaf0f2"
            title_bg = "#07131b"
            title_fg = "#eaf0f2"
            subtitle_fg = "#95adb6"
            section_bg = card_bg
            section_fg = text
            accent_label_bg = card_bg
            accent_label_fg = gold
            muted_bg = card_bg
            muted_fg = text_muted
            appmuted_bg = app_bg
            appmuted_fg = text_muted
            count_bg = "#0f3b38"
            count_fg = "#b2ece2"
            secure_bg = count_bg
            secure_fg = count_fg
            primary_bg = "#2a8e8e"
            primary_fg = "#001414"
            secondary_bg = "#233440"
            secondary_fg = text
            accent_btn_bg = gold
            accent_btn_fg = "#2a1a00"
            field_bg = "#0e1a22"
            field_fg = text
            border = "#2f4851"
            focus = "#2a8e8e"
            tree_bg = card_bg
            tree_field = card_bg
            tree_fg = text
            tree_head_bg = "#2a8e8e"
            tree_head_fg = "#001414"
            tree_sel_bg = "#1d6a6a"
            tree_sel_fg = text
            scroll_bg = "#2f4851"
            scroll_trough = "#0e1a22"
            check_bg = card_bg
            check_fg = text
            gold_bg = "#3d2f00"
            gold_fg = gold
        else:
            app_bg = "#eef3f3"
            card_bg = "#ffffff"
            card_body_bg = "#ffffff"
            sh1 = "#dce4e5"
            sh2 = "#e6ecec"
            text = "#0f2a36"
            text_muted = "#56707a"
            gold = "#d9a13f"
            accent_bg = "#0a2235"
            accent_fg = "#ffffff"
            title_bg = "#0a2235"
            title_fg = "#ffffff"
            subtitle_fg = "#b7d0d8"
            section_bg = card_bg
            section_fg = "#0f2a36"
            accent_label_bg = card_bg
            accent_label_fg = gold
            muted_bg = card_bg
            muted_fg = text_muted
            appmuted_bg = app_bg
            appmuted_fg = text_muted
            count_bg = "#e3f1ee"
            count_fg = "#1c6262"
            secure_bg = count_bg
            secure_fg = count_fg
            primary_bg = "#2a8e8e"
            primary_fg = "#ffffff"
            secondary_bg = "#eaf0f0"
            secondary_fg = "#0f2a36"
            accent_btn_bg = gold
            accent_btn_fg = "#2a1a00"
            field_bg = "#ffffff"
            field_fg = "#0f2a36"
            border = "#cfe1e4"
            focus = "#2a8e8e"
            tree_bg = "#ffffff"
            tree_field = "#ffffff"
            tree_fg = "#0f2a36"
            tree_head_bg = "#1c6262"
            tree_head_fg = "#ffffff"
            tree_sel_bg = "#2a8e8e"
            tree_sel_fg = "#ffffff"
            scroll_bg = "#cfe1e4"
            scroll_trough = "#f2f7f7"
            check_bg = card_bg
            check_fg = "#0f2a36"
            gold_bg = "#fbf0dc"
            gold_fg = "#a8721c"

        style.configure("App.TFrame", background=app_bg)
        style.configure("Card.TFrame", background=card_bg)
        style.configure("CardBody.TFrame", background=card_body_bg)
        style.configure("Shadow1.TFrame", background=sh1)
        style.configure("Shadow2.TFrame", background=sh2)

        style.configure("Title.TLabel", background=title_bg, foreground=title_fg,
                        font=("Segoe UI Semibold", 26, "bold"))
        style.configure("Subtitle.TLabel", background=title_bg, foreground=subtitle_fg,
                        font=("Segoe UI", 10))
        style.configure("Eyebrow.TLabel", background=title_bg, foreground=gold,
                        font=("Segoe UI", 8, "bold"))

        style.configure("Section.TLabel", background=section_bg, foreground=section_fg,
                        font=("Segoe UI Semibold", 13, "bold"))
        style.configure("SectionAccent.TLabel", background=accent_label_bg, foreground=accent_label_fg,
                        font=("Segoe UI", 9, "bold"))
        style.configure("Muted.TLabel", background=muted_bg, foreground=muted_fg,
                        font=("Segoe UI", 9))
        style.configure("AppMuted.TLabel", background=appmuted_bg, foreground=appmuted_fg,
                        font=("Segoe UI", 9))
        style.configure("FieldLabel.TLabel", background=card_bg, foreground=text,
                        font=("Segoe UI Semibold", 9, "bold"))

        style.configure("Count.TLabel", background=count_bg, foreground=count_fg,
                        font=("Segoe UI Semibold", 9, "bold"), padding=(12, 6))
        style.configure("Gold.TLabel", background=gold_bg, foreground=gold_fg,
                        font=("Segoe UI Semibold", 9, "bold"), padding=(12, 6))
        style.configure("Secure.TLabel", background=secure_bg, foreground=secure_fg,
                        font=("Segoe UI Semibold", 9, "bold"), padding=(12, 6))

        style.configure("Primary.TButton", background=primary_bg, foreground=primary_fg,
                        font=("Segoe UI Semibold", 10, "bold"), padding=(18, 11),
                        borderwidth=0, focusthickness=0)
        style.map("Primary.TButton",
                  background=[("active", "#1d6a6a"), ("pressed", "#154d4d")])

        style.configure("Secondary.TButton", background=secondary_bg, foreground=secondary_fg,
                        font=("Segoe UI Semibold", 9, "bold"), padding=(14, 9),
                        borderwidth=0, focusthickness=0)
        style.map("Secondary.TButton",
                  background=[("active", "#d5e0e1" if not dark else "#33495a"),
                              ("pressed", "#c2d2d3" if not dark else "#466175")])

        style.configure("Accent.TButton", background=accent_btn_bg, foreground=accent_btn_fg,
                        font=("Segoe UI Semibold", 9, "bold"), padding=(14, 9),
                        borderwidth=0, focusthickness=0)
        style.map("Accent.TButton",
                  background=[("active", "#c08c2d"), ("pressed", "#a4741d")])

        style.configure("TEntry", fieldbackground=field_bg, foreground=field_fg,
                        bordercolor=border, lightcolor=border, darkcolor=border,
                        padding=9, focusthickness=2, focuscolor=focus,
                        font=("Segoe UI", 10))
        style.map("TEntry",
                  bordercolor=[("focus", focus), ("!focus", border)],
                  lightcolor=[("focus", focus), ("!focus", border)],
                  darkcolor=[("focus", focus), ("!focus", border)])

        style.configure("TCombobox", fieldbackground=field_bg, foreground=field_fg,
                        background=field_bg, arrowcolor=tree_head_fg,
                        bordercolor=border, lightcolor=border, darkcolor=border,
                        padding=8, focusthickness=2, focuscolor=focus,
                        font=("Segoe UI", 10))
        style.map("TCombobox",
                  bordercolor=[("focus", focus), ("!focus", border)],
                  lightcolor=[("focus", focus), ("!focus", border)],
                  darkcolor=[("focus", focus), ("!focus", border)],
                  background=[("readonly", field_bg), ("active", field_bg)],
                  fieldbackground=[("readonly", field_bg)])

        style.configure("Treeview", background=tree_bg, fieldbackground=tree_field,
                        foreground=tree_fg, rowheight=28,
                        bordercolor=border, borderwidth=1, font=("Segoe UI", 9))
        style.configure("Treeview.Heading", background=tree_head_bg, foreground=tree_head_fg,
                        font=("Segoe UI Semibold", 9, "bold"), padding=9,
                        relief="flat", borderwidth=0)
        style.map("Treeview",
                  background=[("selected", tree_sel_bg)],
                  foreground=[("selected", tree_sel_fg)])
        style.map("Treeview.Heading",
                  background=[("active", "#1d6a6a")])

        style.configure("Vertical.TScrollbar", background=scroll_bg, troughcolor=scroll_trough,
                        bordercolor=scroll_bg, arrowcolor=tree_head_fg, arrowsize=14,
                        relief="flat", borderwidth=0, gripcount=0, width=14)
        style.map("Vertical.TScrollbar",
                  background=[("active", focus), ("disabled", scroll_trough)])

        style.configure("Horizontal.TProgressbar",
                        troughcolor=scroll_trough,
                        background=focus,
                        bordercolor=border,
                        lightcolor=focus,
                        darkcolor=focus,
                        thickness=14)

        style.configure("TCheckbutton", background=check_bg, foreground=check_fg,
                        font=("Segoe UI Semibold", 9), focusthickness=0)
        style.map("TCheckbutton",
                  background=[("active", check_bg)],
                  foreground=[("active", check_fg)])

        self._style_colors = {
            "app_bg": app_bg, "card_bg": card_bg, "card_body_bg": card_body_bg,
            "sh1": sh1, "sh2": sh2, "text": text, "text_muted": text_muted,
            "gold": gold, "accent_bg": accent_bg, "title_bg": title_bg,
            "gold_bg": gold_bg, "count_bg": count_bg, "secure_bg": secure_bg,
            "field_bg": field_bg, "border": border, "focus": focus,
        }

    # ---------------------------- Scaffold --------------------------------
    def _build_scaffold(self):
        self._tooltips.clear()
        for child in self.root.winfo_children():
            try:
                child.destroy()
            except Exception:
                pass

    def _install_drain_loop(self):
        self._drain_queue()

    def _drain_queue(self):
        try:
            while True:
                kind, payload = self._queue.get_nowait()
                if kind == "progress":
                    self._set_progress(payload["step"], payload["total"])
                elif kind == "status":
                    self.status.set(str(payload))
                elif kind == "progress_ready_label":
                    self.progress_label.set(self.tr("progress_ready"))
                    if self._progressbar is not None:
                        self._progressbar["value"] = 0
        except queue.Empty:
            pass
        finally:
            self.root.after(80, self._drain_queue)

    def _set_progress(self, step: int, total: int):
        total = max(1, total)
        step = max(0, min(step, total))
        if self._progressbar is not None:
            self._progressbar["maximum"] = total
            self._progressbar["value"] = step
        self.progress_label.set(
            self.tr("progress_label", step=step, total=total,
                    pct=round(step * 100.0 / total, 1))
        )

    def _post(self, kind: str, payload: Any):
        self._queue.put((kind, payload))

    # ---------------------------- Header ----------------------------------
    def _build_header(self):
        dark = self.theme.get() == "dark"
        header_outer = tk.Frame(self.root, bg=self._style_colors["app_bg"], height=160)
        header_outer.pack(fill=X, side="top")
        header_outer.pack_propagate(False)

        title_bg = self._style_colors["title_bg"]
        c = tk.Canvas(header_outer, height=148, highlightthickness=0, bd=0, bg=title_bg)
        c.pack(fill=X, side="top")

        def paint_header(_evt=None):
            w = max(c.winfo_width(), 1)
            c.delete("all")
            if dark:
                self._draw_gradient(c, w, 148, "#07131b", "#0e1a22")
            else:
                self._draw_gradient(c, w, 148, "#0a2235", "#153a52")
            c.create_rectangle(0, 146, w, 148, fill=self._style_colors["gold"], outline="")

        c.bind("<Configure>", paint_header)
        c.after(1, paint_header)

        content = tk.Frame(header_outer, bg=title_bg)
        content.place(x=42, y=28, relwidth=1.0, width=-84)

        ttk.Label(content, text=self.tr("eyebrow"), style="Eyebrow.TLabel").pack(anchor="w")
        ttk.Label(content, text=self.tr("header_title"), style="Title.TLabel").pack(anchor="w", pady=(4, 0))
        ttk.Label(content, text=self.tr("header_subtitle"),
                  style="Subtitle.TLabel").pack(anchor="w", pady=(6, 0))

        controls = tk.Frame(header_outer, bg=title_bg)
        controls.place(relx=1.0, x=-42, y=28, anchor="ne")
        row = tk.Frame(controls, bg=title_bg)
        row.pack(anchor="e")

        tk.Label(row, text=self.tr("lbl_theme") + "  ", bg=title_bg,
                 fg=self._style_colors["text_muted"],
                 font=("Segoe UI Semibold", 9, "bold")).pack(side=LEFT)
        theme_switch = ttk.Combobox(
            row, values=[self.tr("theme_light"), self.tr("theme_dark")],
            state="readonly", width=9
        )
        theme_switch.current(0 if self.theme.get() == "light" else 1)
        theme_switch.pack(side=LEFT, padx=(0, 18))

        tk.Label(row, text=self.tr("lbl_language") + "  ", bg=title_bg,
                 fg=self._style_colors["text_muted"],
                 font=("Segoe UI Semibold", 9, "bold")).pack(side=LEFT)
        langs = available_languages()
        labels = []
        current_idx = 0
        for i, code in enumerate(langs):
            labels.append(code.upper())
            if code == self.language_var.get():
                current_idx = i
        lang_combo = ttk.Combobox(row, values=labels, state="readonly", width=6)
        lang_combo.current(current_idx)
        lang_combo.pack(side=LEFT)

        def on_theme(_e=None):
            chosen = theme_switch.current()
            new_theme = "light" if chosen == 0 else "dark"
            if new_theme == self.theme.get():
                return
            self.theme.set(new_theme)
            self._persist_settings()
            self._rebuild_ui()

        def on_lang(_e=None):
            idx = lang_combo.current()
            if idx < 0 or idx >= len(langs):
                return
            code = langs[idx]
            if code == self.language_code:
                return
            self.language_code = code
            self.language_var.set(code)
            self.language = load_language(code)
            self._persist_settings()
            self._rebuild_ui()

        theme_switch.bind("<<ComboboxSelected>>", on_theme)
        lang_combo.bind("<<ComboboxSelected>>", on_lang)

    def _draw_gradient(self, canvas, w, h, color1, color2):
        steps = max(1, h)
        r1, g1, b1 = canvas.winfo_rgb(color1)
        r2, g2, b2 = canvas.winfo_rgb(color2)
        rr = (r2 - r1) / steps
        rg = (g2 - g1) / steps
        rb = (b2 - b1) / steps
        for i in range(steps):
            nr = int(r1 + rr * i)
            ng = int(g1 + rg * i)
            nb = int(b1 + rb * i)
            color = f"#{nr:04x}{ng:04x}{nb:04x}"
            canvas.create_line(0, i, w, i, fill=color)

    # ----------------------------- Body -----------------------------------
    def _build_body(self):
        outer_wrap = ttk.Frame(self.root, style="App.TFrame")
        outer_wrap.pack(fill=BOTH, expand=True)
        outer_wrap.columnconfigure(0, weight=1)
        outer_wrap.rowconfigure(0, weight=1)

        canvas_wrap = tk.Canvas(outer_wrap, background=self._style_colors["app_bg"],
                                highlightthickness=0, borderwidth=0)
        canvas_wrap.grid(row=0, column=0, sticky="nsew")
        sb = ttk.Scrollbar(outer_wrap, orient="vertical", command=canvas_wrap.yview)
        sb.grid(row=0, column=1, sticky="ns")
        canvas_wrap.configure(yscrollcommand=sb.set)

        scrolled = ttk.Frame(canvas_wrap, style="App.TFrame")
        scrolled_id = canvas_wrap.create_window((0, 0), window=scrolled, anchor="nw")
        self._body_canvas = canvas_wrap
        self._body_scrolled_id = scrolled_id
        self._body_scrolled = scrolled

        def _on_scroll_config(_evt=None):
            canvas_wrap.configure(scrollregion=canvas_wrap.bbox("all"))

        def _on_canvas_config(evt):
            canvas_wrap.itemconfigure(scrolled_id, width=evt.width)

        scrolled.bind("<Configure>", _on_scroll_config)
        canvas_wrap.bind("<Configure>", _on_canvas_config)
        def _on_wheel(evt):
            try:
                if canvas_wrap.winfo_exists():
                    canvas_wrap.yview_scroll(int(-1 * (evt.delta / 120)), "units")
            except Exception:
                pass
        canvas_wrap.bind_all("<MouseWheel>", _on_wheel, add="+")

        self._body = body = ttk.Frame(scrolled, style="App.TFrame", padding=(30, 6, 30, 10))
        body.pack(fill=BOTH, expand=True)
        body.columnconfigure(0, weight=5)
        body.columnconfigure(1, weight=4)
        body.rowconfigure(1, weight=1)
        self._body_layout_root = body

        self._build_card1(body)
        self._build_card2(body)
        self._build_card3(body)
        self._build_footer(body)
        self._apply_breakpoint()

    def _card(self, parent, title=None, subtitle=None, accent=None, **kwargs):
        colors = self._style_colors
        shadow1 = tk.Frame(parent, bg=colors["sh1"], highlightthickness=0)
        shadow2 = tk.Frame(shadow1, bg=colors["sh2"], padx=1, pady=1)
        shadow2.pack(fill=BOTH, expand=True, padx=2, pady=2)
        card = tk.Frame(shadow2, bg=colors["card_bg"], padx=2, pady=2)
        card.pack(fill=BOTH, expand=True, padx=1, pady=1)
        inner = tk.Frame(card, bg=colors["card_body_bg"], padx=26, pady=24)
        inner.pack(fill=BOTH, expand=True)

        head = tk.Frame(inner, bg=colors["card_body_bg"])
        head.pack(fill=X)
        if title or accent:
            left = tk.Frame(head, bg=colors["card_body_bg"])
            left.pack(side=LEFT, fill=X, expand=True)
            if accent:
                ttk.Label(left, text=accent, style="SectionAccent.TLabel").pack(anchor="w")
            if title:
                ttk.Label(left, text=title, style="Section.TLabel").pack(anchor="w", pady=(2, 0))
            tk.Frame(inner, bg=colors["gold"], height=2).pack(fill=X, pady=(16, 22))

        if subtitle:
            ttk.Label(inner, text=subtitle, style="Muted.TLabel").pack(anchor="w", pady=(0, 16))

        body = tk.Frame(inner, bg=colors["card_body_bg"])
        body.pack(fill=BOTH, expand=True)
        return shadow1, body

    def _field_label(self, parent, text, row, col=0, span=2, label_col_width=160):
        lbl = tk.Label(
            parent, text=text, bg=self._style_colors["card_body_bg"],
            fg=self._style_colors["text"],
            font=("Segoe UI Semibold", 9, "bold"), anchor="w",
        )
        lbl.grid(row=row, column=col, sticky="we", padx=(0, 16), pady=(0, 6))
        parent.grid_columnconfigure(col, minsize=label_col_width)
        return lbl

    def _wrap_field(self, parent, row, widget, col=0, span=2, pad_bottom=14):
        container = tk.Frame(parent, bg=self._style_colors["card_body_bg"])
        container.grid(row=row, column=col + 1, sticky="nsew", columnspan=span - 1, pady=(0, pad_bottom))
        container.grid_columnconfigure(0, weight=1)
        widget_master = widget.master
        widget.reparent = None  # noop
        if widget_master is not container:
            widget.pack_forget() if hasattr(widget, "pack_forget") else None
            widget.grid_forget() if hasattr(widget, "grid_forget") else None
        widget.grid(in_=container, row=0, column=0, sticky="ew")
        return container

    # ------------------------- Card 1 (Setup) -----------------------------
    def _build_card1(self, body):
        src_shadow, src_body = self._card(
            body,
            title=self.tr("card01_title"),
            accent=self.tr("card01_accent"),
            subtitle=self.tr("card01_subtitle"),
        )
        src_shadow.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 18))
        src_body.columnconfigure(1, weight=1)
        src_body.columnconfigure(2, weight=0)

        def add_row(r, label_text, var, btn_cmd):
            tk.Label(src_body, text=label_text,
                     bg=self._style_colors["card_body_bg"], fg=self._style_colors["text"],
                     font=("Segoe UI Semibold", 9, "bold"), anchor="w"
                     ).grid(row=r, column=0, sticky="we", padx=(0, 18), pady=(0, 4))
            ent = ttk.Entry(src_body, textvariable=var)
            ent.grid(row=r, column=1, sticky="ew", padx=(0, 10), pady=(0, 14))
            btn = ttk.Button(src_body, text=self.tr("btn_choose_folder"),
                             style="Secondary.TButton", command=btn_cmd)
            btn.grid(row=r, column=2, sticky="ew", pady=(0, 14))
            return ent, btn

        tpl_entry, tpl_btn = add_row(0, self.tr("lbl_templates"), self.template_dir,
                                     self.choose_template_dir)
        out_entry, out_btn = add_row(1, self.tr("lbl_output"), self.output_dir,
                                     self.choose_output_dir)
        self._add_tooltip(tpl_entry, lambda: self.tr("tt_choose_tpl"))
        self._add_tooltip(tpl_btn, lambda: self.tr("tt_choose_tpl"))
        self._add_tooltip(out_entry, lambda: self.tr("tt_choose_out"))
        self._add_tooltip(out_btn, lambda: self.tr("tt_choose_out"))

        action_row = tk.Frame(src_body, bg=self._style_colors["card_body_bg"])
        action_row.grid(row=2, column=0, columnspan=3, sticky="ew", pady=(10, 0))
        refresh_btn = ttk.Button(action_row, text=self.tr("btn_refresh"),
                                 style="Accent.TButton", command=self.refresh_templates)
        refresh_btn.pack(side=LEFT)
        self._add_tooltip(refresh_btn, lambda: self.tr("tt_refresh"))
        dept_btn = ttk.Button(action_row, text=self.tr("btn_depts"),
                              style="Secondary.TButton", command=self.open_department_editor)
        dept_btn.pack(side=LEFT, padx=(10, 0))
        self._add_tooltip(dept_btn, lambda: self.tr("tt_depts"))

    # ------------------------- Card 2 (Employee) --------------------------
    def _build_card2(self, body):
        form_shadow, form_body = self._card(
            body,
            title=self.tr("card02_title"),
            accent=self.tr("card02_accent"),
            subtitle=self.tr("card02_subtitle"),
        )
        self._card2_shadow = form_shadow
        form_shadow.grid(row=1, column=0, sticky="nsew", padx=(0, 18))
        form_body.columnconfigure(1, weight=1)

        self._field_label(form_body, self.tr("lbl_name"), row=0)
        name_entry = ttk.Entry(form_body, textvariable=self.employee_name)
        self._wrap_field(form_body, 0, name_entry)
        self._add_tooltip(name_entry, lambda: self.tr("tt_name"))

        self._field_label(form_body, self.tr("lbl_date"), row=1)
        self.date_picker = DatePickerFrame(form_body, self.language,
                                           bg=self._style_colors["card_body_bg"])
        self.date_picker.grid(row=1, column=1, sticky="nsew", pady=(0, 14))
        self._add_tooltip(self.date_picker, lambda: self.tr("tt_date"))
        for w in (self.date_picker.day_cb, self.date_picker.month_cb, self.date_picker.year_cb):
            self._add_tooltip(w, lambda: self.tr("tt_date"))

        self._field_label(form_body, self.tr("lbl_role"), row=2)
        role_entry = ttk.Entry(form_body, textvariable=self.role)
        self._wrap_field(form_body, 2, role_entry)
        self._add_tooltip(role_entry, lambda: self.tr("tt_role"))

        tk.Label(form_body, text=self.tr("lbl_notes"),
                 bg=self._style_colors["card_body_bg"], fg=self._style_colors["text"],
                 font=("Segoe UI Semibold", 9, "bold"), anchor="w"
                 ).grid(row=3, column=0, sticky="nwe", padx=(0, 16), pady=(0, 6))
        notes_wrap = tk.Frame(form_body, bg=self._style_colors["card_body_bg"])
        notes_wrap.grid(row=3, column=1, sticky="nsew", pady=(0, 14))
        notes_wrap.grid_columnconfigure(0, weight=1)
        colors = self._style_colors
        notes_entry = tk.Text(
            notes_wrap, height=7, wrap="word", font=("Segoe UI", 10),
            bg=colors["field_bg"], fg=colors["text"], relief="solid", borderwidth=1,
            highlightthickness=2, highlightbackground=colors["border"],
            highlightcolor=colors["focus"], padx=10, pady=9,
            insertbackground=colors["focus"],
        )
        notes_entry.grid(row=0, column=0, sticky="nsew")
        self.notes_widget = notes_entry
        sb_notes = ttk.Scrollbar(notes_wrap, orient="vertical", command=notes_entry.yview)
        sb_notes.grid(row=0, column=1, sticky="ns")
        notes_entry.configure(yscrollcommand=sb_notes.set)
        self._add_tooltip(notes_entry, lambda: self.tr("tt_notes"))
        form_body.rowconfigure(3, weight=1)

        bottom_form = tk.Frame(form_body, bg=self._style_colors["card_body_bg"])
        bottom_form.grid(row=4, column=0, columnspan=2, sticky="ew", pady=(6, 0))
        auto_cb = ttk.Checkbutton(bottom_form, text=self.tr("cb_autoopen"),
                                  variable=self.auto_open)
        auto_cb.pack(side=LEFT)
        self._add_tooltip(auto_cb, lambda: self.tr("tt_autoopen"))

    # ----------------------- Card 3 (Summary) -----------------------------
    def _build_card3(self, body):
        prev_shadow, prev_body = self._card(
            body,
            title=self.tr("card03_title"),
            accent=self.tr("card03_accent"),
            subtitle=self.tr("card03_subtitle"),
        )
        self._card3_shadow = prev_shadow
        prev_shadow.grid(row=1, column=1, sticky="nsew")
        prev_body.columnconfigure(0, weight=1)

        # --- Reparto selector ---
        dept_frame = tk.Frame(prev_body, bg=self._style_colors["card_body_bg"])
        dept_frame.grid(row=0, column=0, sticky="ew", pady=(0, 14))
        dept_frame.columnconfigure(1, weight=1)

        multi_cb = ttk.Checkbutton(dept_frame, text=self.tr("cb_multidept"),
                                   variable=self.multi_dept_mode,
                                   command=self._toggle_multi_dept)
        multi_cb.grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 6))
        self._add_tooltip(multi_cb, lambda: self.tr("tt_multidept"))

        tk.Label(dept_frame, text=self.tr("lbl_department"),
                 bg=self._style_colors["card_body_bg"], fg=self._style_colors["text"],
                 font=("Segoe UI Semibold", 9, "bold"), anchor="w"
                 ).grid(row=1, column=0, sticky="w", padx=(0, 14), pady=(0, 6))

        single_wrap = tk.Frame(dept_frame, bg=self._style_colors["card_body_bg"])
        single_wrap.grid(row=1, column=1, sticky="ew", pady=(0, 6))
        single_wrap.columnconfigure(0, weight=1)
        self.department_combo = ttk.Combobox(
            single_wrap, state="readonly", textvariable=self.department,
            font=("Segoe UI Semibold", 10, "bold"), height=14,
        )
        self.department_combo.grid(row=0, column=0, sticky="ew")
        self.department_combo.bind("<<ComboboxSelected>>", lambda _e: self.update_document_list())
        self._single_dept_wrap = single_wrap

        multi_wrap = tk.Frame(dept_frame, bg=self._style_colors["card_body_bg"])
        multi_wrap.grid(row=2, column=0, columnspan=2, sticky="ew")
        multi_wrap.grid_remove()
        self._multi_dept_wrap = multi_wrap
        self._render_multi_dept_list(multi_wrap)

        tk.Label(dept_frame, text="  " + self.tr("lbl_department_hint").lstrip(),
                 bg=self._style_colors["card_body_bg"], fg=self._style_colors["text_muted"],
                 font=("Segoe UI", 8), anchor="w",
                 ).grid(row=3, column=0, columnspan=2, sticky="w")

        self._add_tooltip(self.department_combo, lambda: self.tr("tt_dept"))

        # --- Badge count + select all/none ---
        badge_row = tk.Frame(prev_body, bg=self._style_colors["card_body_bg"])
        badge_row.grid(row=1, column=0, sticky="ew", pady=(0, 12))
        ttk.Label(badge_row, textvariable=self.count_label, style="Count.TLabel").pack(side=LEFT)
        right_badges = tk.Frame(badge_row, bg=self._style_colors["card_body_bg"])
        right_badges.pack(side=RIGHT)
        btn_all = ttk.Button(right_badges, text=self.tr("btn_selectall"),
                             style="Secondary.TButton",
                             command=lambda: self._set_all_inclusion(True))
        btn_none = ttk.Button(right_badges, text=self.tr("btn_selectnone"),
                              style="Secondary.TButton",
                              command=lambda: self._set_all_inclusion(False))
        btn_all.pack(side=LEFT, padx=(0, 6))
        btn_none.pack(side=LEFT)
        self._add_tooltip(btn_all, lambda: self.tr("tt_include"))
        self._add_tooltip(btn_none, lambda: self.tr("tt_include"))

        # --- Treeview with checkboxes ---
        tree_wrap = tk.Frame(prev_body, bg=self._style_colors["card_body_bg"])
        tree_wrap.grid(row=2, column=0, sticky="nsew", pady=(0, 4))
        tree_wrap.rowconfigure(0, weight=1)
        tree_wrap.columnconfigure(0, weight=1)
        prev_body.rowconfigure(2, weight=1)

        self.tree = ttk.Treeview(
            tree_wrap,
            columns=("include", "documento", "copie", "stato"),
            show="headings",
            height=10,
            selectmode="none",
        )
        self.tree.heading("include", text=self.tr("col_include"))
        self.tree.heading("documento", text=self.tr("col_document"))
        self.tree.heading("copie", text=self.tr("col_copies"))
        self.tree.heading("stato", text=self.tr("col_status"))
        self.tree.column("include", width=46, anchor="center", stretch=False)
        self.tree.column("documento", width=280, anchor="w", stretch=True)
        self.tree.column("copie", width=70, anchor="center", stretch=False)
        self.tree.column("stato", width=96, anchor="center", stretch=False)
        self.tree.grid(row=0, column=0, sticky="nsew")
        tree_sb = ttk.Scrollbar(tree_wrap, orient="vertical", command=self.tree.yview)
        tree_sb.grid(row=0, column=1, sticky="ns")
        self.tree.configure(yscrollcommand=tree_sb.set)
        self.tree.bind("<Button-1>", self._on_tree_click)
        self.tree.bind("<space>", lambda _e: self._toggle_focused_row())
        self._add_tooltip(self.tree, lambda: self.tr("tt_include"))

        # --- Colors hash status tags ---
        colors = self._style_colors
        self.tree.tag_configure("odd", background=colors["card_body_bg"])
        self.tree.tag_configure("even", background="#f3faf8" if self.theme.get() == "light" else "#1a2e38")
        self.tree.tag_configure("tutti", background="#fff8ea" if self.theme.get() == "light" else "#3a2e0a")
        self.tree.tag_configure("modified", background="#fff5b3" if self.theme.get() == "light" else "#554500")
        self.tree.tag_configure("new", background="#d4f5d4" if self.theme.get() == "light" else "#103b1f")

        # --- Progress + buttons ---
        prog_wrap = tk.Frame(prev_body, bg=self._style_colors["card_body_bg"])
        prog_wrap.grid(row=3, column=0, sticky="ew", pady=(4, 0))
        self._progressbar = ttk.Progressbar(prog_wrap, orient="horizontal",
                                             mode="determinate", maximum=100, value=0)
        self._progressbar.pack(fill=X, side="top")
        ttk.Label(prog_wrap, textvariable=self.progress_label, style="Muted.TLabel"
                  ).pack(anchor="w", side="top", pady=(4, 0))

        gen_frame = tk.Frame(prev_body, bg=self._style_colors["card_body_bg"])
        gen_frame.grid(row=4, column=0, sticky="ew", pady=(14, 0))
        gen_btn = ttk.Button(gen_frame, text=self.tr("btn_generate"),
                             style="Primary.TButton", command=self.generate)
        gen_btn.pack(side=RIGHT, ipadx=18, ipady=5)
        batch_btn = ttk.Button(gen_frame, text=self.tr("btn_batch"),
                               style="Secondary.TButton", command=self.open_batch_window)
        batch_btn.pack(side=RIGHT, padx=(0, 10), ipady=5)
        tk.Label(gen_frame, text=self.tr("lbl_generate_hint"),
                 bg=self._style_colors["card_body_bg"], fg=self._style_colors["text_muted"],
                 font=("Segoe UI", 8), anchor="w",
                 ).pack(side=LEFT, padx=(4, 0))
        self._add_tooltip(gen_btn, lambda: self.tr("tt_generate"))
        self._add_tooltip(batch_btn, lambda: self.tr("tt_batch"))

        help_frame = tk.Frame(prev_body, bg=self._style_colors["card_body_bg"])
        help_frame.grid(row=5, column=0, sticky="ew", pady=(14, 0))
        tk.Frame(help_frame, bg=self._style_colors["count_bg"], width=4, height=56).pack(side=LEFT)
        help_wrap_length = 380
        tip = tk.Label(
            help_frame, text=self.tr("help_tip_body"),
            bg=self._style_colors["card_body_bg"], fg=self._style_colors["text_muted"],
            font=("Segoe UI", 9), justify="left", anchor="w",
            padx=12, pady=8, wraplength=help_wrap_length,
        )
        self._help_tip_label = tip
        tip.pack(side=LEFT, fill=X, expand=True)

        self._toggle_multi_dept()

    # ---------------------- Footer ----------------------------------------
    def _build_footer(self, body_container):
        footer_outer = tk.Frame(self.root, bg=self._style_colors["app_bg"])
        footer_outer.pack(fill=X, side="bottom")
        footer = tk.Frame(footer_outer, bg=self._style_colors["card_body_bg"], padx=24, pady=12)
        footer.pack(fill=X, padx=30, pady=(0, 16))

        status_wrap = tk.Frame(footer, bg=self._style_colors["card_body_bg"])
        status_wrap.pack(side=LEFT, fill=X, expand=True)
        tk.Label(status_wrap, text="●", bg=self._style_colors["card_body_bg"],
                 fg=self._style_colors["focus"], font=("Segoe UI", 10, "bold"),
                 ).pack(side=LEFT)
        tk.Label(status_wrap, textvariable=self.status,
                 bg=self._style_colors["card_body_bg"], fg=self._style_colors["text"],
                 font=("Segoe UI Semibold", 9), anchor="w", padx=8,
                 ).pack(side=LEFT, fill=X, expand=True)
        tk.Label(status_wrap, textvariable=self.hash_stat_label,
                 bg=self._style_colors["card_body_bg"], fg=self._style_colors["text_muted"],
                 font=("Segoe UI", 8), anchor="w", padx=14,
                 ).pack(side=LEFT)

        secure = ttk.Label(footer, text=self.tr("secure_label"), style="Secure.TLabel")
        secure.pack(side=RIGHT)
        self._footer = footer_outer

    # ---------------------- Tooltip helper --------------------------------
    def _add_tooltip(self, widget, getter):
        self._tooltips.append(Tooltip(widget, getter))

    # ---------------------- Multi reparto UI toggle ----------------------
    def _render_multi_dept_list(self, parent):
        for child in parent.winfo_children():
            child.destroy()
        self.multi_dept_values.clear()
        options = department_options(self.templates)
        if not options:
            tk.Label(parent, text="—", bg=self._style_colors["card_body_bg"],
                     fg=self._style_colors["text_muted"]).pack(anchor="w")
            return
        cols = min(3, max(1, (len(options) + 3) // 4))
        for i, dept in enumerate(options):
            var = BooleanVar(value=False)
            self.multi_dept_values[dept] = var
            cb = ttk.Checkbutton(parent, text=dept, variable=var,
                                 command=self.update_document_list)
            cb.grid(row=i // cols, column=i % cols, sticky="w", padx=(0, 10), pady=2)
            self._add_tooltip(cb, lambda d=dept: self.tr("tt_dept") + f" [{d}]")

    def _toggle_multi_dept(self):
        multi = self.multi_dept_mode.get()
        if multi:
            self._single_dept_wrap.grid_remove()
            self._multi_dept_wrap.grid()
            self._render_multi_dept_list(self._multi_dept_wrap)
            self.department.set("")
        else:
            self._multi_dept_wrap.grid_remove()
            self._single_dept_wrap.grid(row=1, column=1, sticky="ew", pady=(0, 6))
        self.update_document_list()

    def _current_departments(self) -> list[str]:
        if self.multi_dept_mode.get():
            return [d for d, v in self.multi_dept_values.items() if v.get()]
        d = self.department.get().strip().upper()
        return [d] if d else []

    # ---------------------- Treeview inclusion logic ----------------------
    def _current_selected_templates(self) -> list[TemplateFile]:
        depts = self._current_departments()
        if not depts:
            return []
        if len(depts) == 1:
            chosen = templates_for_department(self.templates, depts[0])
        else:
            chosen = templates_for_departments(self.templates, depts)
        return [t for t in chosen if self.template_inclusion.get(t.path, True)]

    def _set_all_inclusion(self, value: bool):
        for path in list(self.template_inclusion.keys()):
            self.template_inclusion[path] = value
        self.update_document_list()

    def _on_tree_click(self, event):
        region = self.tree.identify("region", event.x, event.y)
        col = self.tree.identify_column(event.x)
        item = self.tree.identify_row(event.y)
        if not item:
            return
        if region == "cell" and col == "#1":
            self._toggle_row(item)
            return
        if region == "heading" and col == "#1":
            current_all = all(
                self.template_inclusion.get(self._row_path.get(item), True)
                for item in self.tree.get_children()
            ) if self.tree.get_children() else True
            self._set_all_inclusion(not current_all)

    def _toggle_focused_row(self):
        item = self.tree.focus()
        if item:
            self._toggle_row(item)

    def _toggle_row(self, item):
        path = self._row_path.get(item)
        if path is None:
            return
        self.template_inclusion[path] = not self.template_inclusion.get(path, True)
        self.update_document_list()

    # ---------------------- Breakpoint responsive ------------------------
    def _on_root_resize(self, evt):
        if evt.widget is not self.root:
            return
        self._current_width = evt.width
        try:
            self._apply_breakpoint()
        except Exception:
            pass

    def _apply_breakpoint(self):
        want = "narrow" if self._current_width and self._current_width < 1100 else "wide"
        if want == self._current_breakpoint:
            return
        self._current_breakpoint = want
        layout = getattr(self, "_body_layout_root", None)
        if layout is None:
            return
        if want == "narrow":
            if self._card2_shadow is not None:
                self._card2_shadow.grid_configure(row=1, column=0, columnspan=2,
                                                   sticky="nsew", padx=(0, 0), pady=(0, 18))
            if self._card3_shadow is not None:
                self._card3_shadow.grid_configure(row=2, column=0, columnspan=2, sticky="nsew")
            layout.rowconfigure(1, weight=1)
            layout.rowconfigure(2, weight=1)
        else:
            if self._card2_shadow is not None:
                self._card2_shadow.grid_configure(row=1, column=0, columnspan=1,
                                                   sticky="nsew", padx=(0, 18), pady=(0, 0))
            if self._card3_shadow is not None:
                self._card3_shadow.grid_configure(row=1, column=1, columnspan=1, sticky="nsew")
            layout.rowconfigure(2, weight=0)

    # ---------------------- Rebuild on theme/lang change ------------------
    def _rebuild_ui(self):
        self._configure_style()
        self._apply_theme_root()
        self.root.title(self.tr("app_title"))
        # Save transient non-StringVar state before destroy
        saved_notes = ""
        saved_date = None
        try:
            if hasattr(self, "notes_widget") and self.notes_widget is not None:
                if self.notes_widget.winfo_exists():
                    saved_notes = self.notes_widget.get("1.0", "end-1c")
        except Exception:
            saved_notes = ""
        try:
            if hasattr(self, "date_picker") and self.date_picker is not None:
                try:
                    if self.date_picker.winfo_exists():
                        saved_date = self.date_picker.get_date()
                except Exception:
                    saved_date = None
        except Exception:
            saved_date = None
        # Status/count/progress StringVar: reapply keys that are language-dependent
        self.status.set(self.tr("status_initial"))
        self.count_label.set(self.tr("count_none"))
        self.progress_label.set(self.tr("progress_ready"))
        for child in list(self.root.winfo_children()):
            try:
                child.destroy()
            except Exception:
                pass
        self._build_header()
        self._build_body()
        self.refresh_templates()
        # Restore transient non-StringVar state after rebuild
        try:
            if saved_notes and hasattr(self, "notes_widget") and self.notes_widget is not None:
                try:
                    if self.notes_widget.winfo_exists():
                        self.notes_widget.delete("1.0", "end")
                        self.notes_widget.insert("1.0", saved_notes)
                except Exception:
                    pass
        except Exception:
            pass
        try:
            if saved_date is not None and hasattr(self, "date_picker") and self.date_picker is not None:
                try:
                    if self.date_picker.winfo_exists():
                        self.date_picker.set_date(saved_date)
                        self.date_picker.configure_language(self.language)
                except Exception:
                    pass
        except Exception:
            pass

    # ---------------------- Dir choices -----------------------------------
    def choose_template_dir(self) -> None:
        chosen = filedialog.askdirectory(initialdir=self.template_dir.get(),
                                         title=self.tr("btn_choose_folder"))
        if chosen:
            self.template_dir.set(chosen)
            self._persist_settings()
            self.refresh_templates()

    def choose_output_dir(self) -> None:
        chosen = filedialog.askdirectory(initialdir=self.output_dir.get(),
                                         title=self.tr("btn_choose_folder"))
        if chosen:
            self.output_dir.set(chosen)
            self._persist_settings()

    # ---------------------- Refresh templates -----------------------------
    def refresh_templates(self) -> None:
        folder = Path(self.template_dir.get()).expanduser()
        self.templates, self.ignored = discover_templates(folder)
        departments = department_options(self.templates)
        self.department_combo["values"] = departments
        if departments and self.department.get().upper() not in departments:
            self.department.set(departments[0])
        elif not departments:
            self.department.set("")
        if hasattr(self, "_multi_dept_wrap") and self._multi_dept_wrap is not None:
            self._render_multi_dept_list(self._multi_dept_wrap)
        for tpl in self.templates:
            if tpl.path not in self.template_inclusion:
                self.template_inclusion[tpl.path] = True
        stale = [p for p in list(self.template_inclusion.keys())
                 if not any(t.path == p for t in self.templates)]
        for s in stale:
            del self.template_inclusion[s]
        self.hash_status = classify_template_hashes(self.templates, self.saved_hashes)
        self.update_document_list()
        ok = sum(1 for v in self.hash_status.values() if v == "ok")
        mod = sum(1 for v in self.hash_status.values() if v == "modified")
        new = sum(1 for v in self.hash_status.values() if v == "new")
        self.hash_stat_label.set(self.tr("hash_status", ok=ok, mod=mod, new=new))
        if not folder.exists():
            self.status.set("La cartella template non esiste ancora.")
        elif not self.templates:
            self.status.set("Nessun template valido. Usa REPARTO_NUMERO_CODICE.")
        elif self.ignored:
            self.status.set(f"{len(self.templates)} template; {len(self.ignored)} ignorati.")
        else:
            self.status.set(f"{len(self.templates)} template pronti.")

    def update_document_list(self) -> None:
        if not hasattr(self, "tree") or self.tree is None:
            return
        for item in self.tree.get_children():
            self.tree.delete(item)
        self._row_path: dict[str, Path] = {}
        depts = self._current_departments()
        if len(depts) == 1:
            selected = templates_for_department(self.templates, depts[0])
        elif depts:
            selected = templates_for_departments(self.templates, depts)
        else:
            selected = []
        total_copies = sum(template.copies for template in selected
                           if self.template_inclusion.get(template.path, True))
        total_templates = sum(1 for template in selected
                              if self.template_inclusion.get(template.path, True))
        self.count_label.set(self.tr("count_label", total=total_copies, templates=total_templates))
        for idx, template in enumerate(selected):
            scope = "Tutti" if template.is_for_every_department else template.department.upper()
            included = self.template_inclusion.get(template.path, True)
            mark = "☑" if included else "☐"
            status_key = self.hash_status.get(str(template.path), "ok")
            if status_key == "modified":
                status_text = self.tr("status_modified")
            elif status_key == "new":
                status_text = self.tr("status_new")
            else:
                status_text = self.tr("status_ok")
            tags: list[str] = []
            if template.is_for_every_department:
                tags.append("tutti")
            elif idx % 2 == 0:
                tags.append("even")
            else:
                tags.append("odd")
            if status_key in {"modified", "new"}:
                tags.append(status_key)
            item = self.tree.insert(
                "", END,
                values=(mark, f"  {scope} · {template.path.name}",
                        template.copies, status_text),
                tags=tuple(tags),
            )
            self._row_path[item] = template.path

    # ---------------------- Generate single -------------------------------
    def generate(self) -> None:
        if self._worker_active:
            return
        name = self.employee_name.get().strip()
        entry_date = self.date_picker.get_string()
        departments = self._current_departments()
        if not name:
            messagebox.showwarning(self.tr("mb_missing_title"), self.tr("mb_missing_name"))
            return
        if not entry_date:
            messagebox.showwarning(self.tr("mb_missing_title"), self.tr("mb_missing_date"))
            return
        if not departments:
            messagebox.showwarning(self.tr("mb_missing_title"), self.tr("mb_missing_dept"))
            return
        selected = self._current_selected_templates()
        if not selected:
            messagebox.showwarning(self.tr("mb_no_docs_title"), self.tr("mb_no_docs_body"))
            return
        output_dir_path = Path(self.output_dir.get()).expanduser()
        dept_tag = safe_file_part("+".join(departments))
        base = output_dir_path / f"dossier_{safe_file_part(name)}_{dept_tag}.pdf"
        output_path = base
        counter = 2
        while output_path.exists():
            output_path = output_dir_path / (
                f"dossier_{safe_file_part(name)}_{dept_tag}_{counter}.pdf"
            )
            counter += 1
        notes = self.notes_widget.get("1.0", END).strip()
        dept_str = "+".join(departments)
        self._worker_active = True
        self.status.set(self.tr("mb_progress_status"))
        self._post("progress_ready_label", None)
        self._persist_settings()
        # --- Read ALL Tk/StringVar values IN MAIN THREAD before starting worker ---
        role_value = self.role.get().strip()
        auto_open_value = bool(self.auto_open.get())
        saved_snapshot = dict(self.saved_hashes) if self.saved_hashes else {}
        tr_done = self.tr("mb_status_done", name=output_path.name)
        tr_error = self.tr("mb_status_error")

        def progress_cb(step: int, total: int):
            self._post("progress", {"step": step, "total": total})

        def work():
            try:
                total = build_pdf(output_path, name, entry_date, dept_str,
                                  role_value, notes, selected,
                                  progress_cb=progress_cb)
                self._save_history(output_path, name, dept_str, total)
                hashes = dict(saved_snapshot)
                for t in selected:
                    try:
                        hashes[str(t.path)] = compute_template_hash(t.path)
                    except Exception:
                        pass
                save_hashes(hashes)
                self.saved_hashes = hashes
                self._post("status", tr_done)
                self._post("done_single", {
                    "ok": True, "path": str(output_path), "total": total,
                    "auto_open": auto_open_value,
                })
            except Exception as error:  # noqa: BLE001
                traceback.print_exc()
                self._post("status", tr_error)
                self._post("done_single", {"ok": False, "error": str(error)})
            finally:
                self._post("worker_done", None)

        threading.Thread(target=work, daemon=True).start()

    # ---------------------- Worker -> mainloop routing --------------------
    def _drain_queue(self):  # noqa: F811 (intended override; no, we already defined it. Actually duplicate name! Fix: call earlier install_drain_loop -> _drain_queue, but this method will conflict. Fix by renaming the worker drain differently: actually previous code already installed _drain_queue but we redefined it here → conflict. Fix: merge the two handlers.
        try:
            while True:
                kind, payload = self._queue.get_nowait()
                if kind == "progress":
                    self._set_progress(payload["step"], payload["total"])
                elif kind == "status":
                    self.status.set(str(payload))
                elif kind == "progress_ready_label":
                    self.progress_label.set(self.tr("progress_ready"))
                    if self._progressbar is not None:
                        self._progressbar["value"] = 0
                elif kind == "done_single":
                    if payload.get("ok"):
                        messagebox.showinfo(
                            self.tr("mb_done_title"),
                            self.tr("mb_done_body", total=payload.get("total", 0),
                                    path=payload.get("path", "")),
                        )
                        if payload.get("auto_open"):
                            open_folder(Path(str(payload["path"])).parent)
                        self.refresh_templates()
                    else:
                        messagebox.showerror(self.tr("mb_error_title"), str(payload.get("error", "")))
                elif kind == "done_batch":
                    summary = (
                        f"{payload.get('ok', 0)} " + self.tr("bat_summary_ok") +
                        "   ·   " + f"{payload.get('skip', 0)} " + self.tr("bat_summary_skip") +
                        "   ·   " + f"{payload.get('fail', 0)} " + self.tr("bat_summary_fail") +
                        "\n\n" + str(payload.get("detail", ""))
                    )
                    if payload.get("auto_open") and payload.get("out_dir"):
                        open_folder(Path(str(payload["out_dir"])))
                    messagebox.showinfo(self.tr("bat_summary_title"), summary)
                    self.refresh_templates()
                elif kind == "worker_done":
                    self._worker_active = False
                    self.saved_hashes = load_saved_hashes()
                    self.hash_status = classify_template_hashes(self.templates, self.saved_hashes)
                    ok = sum(1 for v in self.hash_status.values() if v == "ok")
                    mod = sum(1 for v in self.hash_status.values() if v == "modified")
                    new = sum(1 for v in self.hash_status.values() if v == "new")
                    self.hash_stat_label.set(self.tr("hash_status", ok=ok, mod=mod, new=new))
        except queue.Empty:
            pass
        finally:
            self.root.after(80, self._drain_queue)

    def _install_drain_loop(self):
        self.root.after(80, self._drain_queue)

    # ---------------------- Department Editor -----------------------------
    def open_department_editor(self):
        win = tk.Toplevel(self.root)
        win.title(self.tr("de_title"))
        win.transient(self.root)
        win.grab_set()
        win.geometry("520x480")
        win.configure(bg=self._style_colors["app_bg"])
        body = tk.Frame(win, bg=self._style_colors["card_body_bg"], padx=18, pady=18)
        body.pack(fill=BOTH, expand=True, padx=14, pady=14)
        ttk.Label(body, text=self.tr("de_hint"), style="Muted.TLabel").pack(anchor="w", pady=(0, 12))
        list_wrap = tk.Frame(body, bg=self._style_colors["card_body_bg"])
        list_wrap.pack(fill=BOTH, expand=True, pady=(0, 12))
        list_wrap.columnconfigure(0, weight=1)
        list_wrap.rowconfigure(0, weight=1)
        lb = tk.Listbox(list_wrap, activestyle="dotbox",
                        bg=self._style_colors["field_bg"], fg=self._style_colors["text"],
                        font=("Segoe UI", 10), selectmode="single",
                        highlightthickness=2,
                        highlightbackground=self._style_colors["border"],
                        highlightcolor=self._style_colors["focus"],
                        relief="flat", borderwidth=0)
        lb.grid(row=0, column=0, sticky="nsew")
        sb = ttk.Scrollbar(list_wrap, orient="vertical", command=lb.yview)
        sb.grid(row=0, column=1, sticky="ns")
        lb.configure(yscrollcommand=sb.set)
        current = load_departments_from_file()
        for item in current:
            lb.insert(END, item)
        row = tk.Frame(body, bg=self._style_colors["card_body_bg"])
        row.pack(fill=X, pady=(0, 12))
        tk.Label(row, text=self.tr("de_entry"), bg=self._style_colors["card_body_bg"],
                 fg=self._style_colors["text_muted"],
                 font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 4))
        entry = ttk.Entry(row)
        entry.pack(fill=X, pady=(0, 10))
        btns = tk.Frame(row, bg=self._style_colors["card_body_bg"])
        btns.pack(fill=X)

        def add():
            val = entry.get().strip().upper()
            if not val:
                messagebox.showwarning(self.tr("de_err_empty_title"), self.tr("de_err_empty_body"))
                return
            if val in lb.get(0, END):
                messagebox.showwarning(self.tr("de_err_exists_title"), self.tr("de_err_exists_body"))
                return
            lb.insert(END, val)
            entry.delete(0, END)

        def rename():
            sel = lb.curselection()
            if not sel:
                messagebox.showwarning(self.tr("de_err_none_title"), self.tr("de_err_none_body"))
                return
            new_val = entry.get().strip().upper()
            if not new_val:
                messagebox.showwarning(self.tr("de_err_empty_title"), self.tr("de_err_empty_body"))
                return
            if new_val in lb.get(0, END):
                messagebox.showwarning(self.tr("de_err_exists_title"), self.tr("de_err_exists_body"))
                return
            lb.delete(sel[0])
            lb.insert(sel[0], new_val)

        def delete():
            sel = lb.curselection()
            if not sel:
                messagebox.showwarning(self.tr("de_err_none_title"), self.tr("de_err_none_body"))
                return
            lb.delete(sel[0])

        ttk.Button(btns, text=self.tr("de_add"), style="Secondary.TButton", command=add
                   ).pack(side=LEFT, padx=(0, 8))
        ttk.Button(btns, text=self.tr("de_rename"), style="Secondary.TButton", command=rename
                   ).pack(side=LEFT, padx=(0, 8))
        ttk.Button(btns, text=self.tr("de_delete"), style="Secondary.TButton", command=delete
                   ).pack(side=LEFT)

        footer = tk.Frame(body, bg=self._style_colors["card_body_bg"])
        footer.pack(fill=X, side="bottom")

        def save_and_close():
            new_list = list(dict.fromkeys(lb.get(0, END)))  # preserve order, unique
            try:
                DEPARTMENTS_FILE.parent.mkdir(parents=True, exist_ok=True)
                content = "\n".join(new_list) + ("\n" if new_list else "")
                DEPARTMENTS_FILE.write_text(content, encoding="utf-8")
            except OSError as exc:
                messagebox.showerror(self.tr("mb_error_title"), str(exc))
                return
            clear_caches()
            self.refresh_templates()
            win.destroy()

        ttk.Button(footer, text=self.tr("de_save"), style="Primary.TButton",
                   command=save_and_close).pack(side=RIGHT)
        ttk.Button(footer, text=self.tr("de_cancel"), style="Secondary.TButton",
                   command=win.destroy).pack(side=RIGHT, padx=(0, 10))

    # ----------------------- Batch Window ----------------------------------
    def open_batch_window(self):
        win = tk.Toplevel(self.root)
        win.title(self.tr("bat_title"))
        win.transient(self.root)
        win.grab_set()
        win.geometry("760x560")
        win.configure(bg=self._style_colors["app_bg"])
        body = tk.Frame(win, bg=self._style_colors["card_body_bg"], padx=18, pady=18)
        body.pack(fill=BOTH, expand=True, padx=14, pady=14)
        ttk.Label(body, text=self.tr("bat_subtitle"), style="Muted.TLabel").pack(anchor="w", pady=(0, 12))
        path_var = StringVar()

        top_row = tk.Frame(body, bg=self._style_colors["card_body_bg"])
        top_row.pack(fill=X, pady=(0, 10))
        ttk.Entry(top_row, textvariable=path_var).pack(side=LEFT, fill=X, expand=True, padx=(0, 10))
        tree_holder = tk.Frame(body, bg=self._style_colors["card_body_bg"])
        tree_holder.pack(fill=BOTH, expand=True, pady=(0, 12))
        tree_holder.rowconfigure(0, weight=1)
        tree_holder.columnconfigure(0, weight=1)
        pv = ttk.Treeview(tree_holder, columns=("nome", "data", "reparto", "ruolo", "note"),
                          show="headings", height=14)
        for col, title in (("nome", "Nome"), ("data", "Data"), ("reparto", "Reparto"),
                           ("ruolo", "Ruolo"), ("note", "Note")):
            pv.heading(col, text=title)
            pv.column(col, width=140 if col in {"nome", "ruolo", "note"} else 110, anchor="w")
        pv.grid(row=0, column=0, sticky="nsew")
        vsb = ttk.Scrollbar(tree_holder, orient="vertical", command=pv.yview)
        vsb.grid(row=0, column=1, sticky="ns")
        pv.configure(yscrollcommand=vsb.set)

        rows_data: list[dict[str, str]] = []

        def load():
            p = filedialog.askopenfilename(
                title=self.tr("bat_choose"),
                initialdir=self.template_dir.get(),
                filetypes=[("CSV / Excel", "*.csv *.xlsx *.xls"), ("All", "*.*")],
            )
            if not p:
                return
            path_var.set(p)
            try:
                loaded = self._parse_batch_file(Path(p))
            except Exception as exc:  # noqa: BLE001
                messagebox.showerror(self.tr("mb_error_title"), str(exc))
                return
            nonlocal rows_data
            rows_data = loaded
            for item in pv.get_children():
                pv.delete(item)
            for r in rows_data:
                pv.insert("", END, values=(r.get("Nome", ""), r.get("Data", ""),
                                           r.get("Reparto", ""), r.get("Ruolo", ""), r.get("Note", "")))

        def run():
            if not rows_data:
                return
            if self._worker_active:
                return
            out_dir_path = Path(self.output_dir.get()).expanduser()
            self._worker_active = True
            self._post("progress_ready_label", None)
            # --- Snapshot ALL Tk/self state in MAIN THREAD ---
            auto_open_batch = bool(self.auto_open.get())
            lang_snap = dict(self.language) if self.language else {}
            def s_tr(key: str, **kw) -> str:
                raw = lang_snap.get(key, key)
                try:
                    return raw.format(**kw) if kw else raw
                except Exception:
                    return raw
            tpl_snap = list(self.templates)
            inc_snap = dict(self.template_inclusion)
            saved_snap = dict(self.saved_hashes) if self.saved_hashes else {}
            dept_opts = {d.upper() for d in department_options(tpl_snap)}
            out_dir_snap = Path(out_dir_path)

            def work2():
                ok = skip = fail = 0
                details: list[str] = []
                total = max(1, len(rows_data))
                try:
                    for idx, row in enumerate(rows_data, start=1):
                        self._post("progress", {"step": idx, "total": total})
                        nome = (row.get("Nome") or row.get("name") or "").strip()
                        data = (row.get("Data") or row.get("date") or "").strip()
                        reparto = (row.get("Reparto") or row.get("department") or "").strip()
                        ruolo = (row.get("Ruolo") or row.get("role") or "").strip()
                        note = (row.get("Note") or row.get("notes") or "").strip()
                        if not nome:
                            skip += 1
                            details.append(s_tr("bat_skip_empty", i=idx))
                            continue
                        parsed = self._parse_date_str(data)
                        if parsed is None:
                            skip += 1
                            details.append(s_tr("bat_skip_date", i=idx))
                            continue
                        dept_ok = reparto.upper() in dept_opts
                        if not dept_ok:
                            skip += 1
                            details.append(s_tr("bat_skip_dept", i=idx, d=reparto))
                            continue
                        sel = templates_for_department(tpl_snap, reparto)
                        filt = [t for t in sel if inc_snap.get(t.path, True)]
                        if not filt:
                            skip += 1
                            details.append(s_tr("bat_skip_dept", i=idx, d=reparto))
                            continue
                        out_file = out_dir_snap / (
                            f"dossier_{safe_file_part(nome)}_{safe_file_part(reparto)}.pdf"
                        )
                        c = 2
                        while out_file.exists():
                            out_file = out_dir_snap / (
                                f"dossier_{safe_file_part(nome)}_{safe_file_part(reparto)}_{c}.pdf"
                            )
                            c += 1
                        try:
                            build_pdf(out_file, nome, parsed.strftime("%d/%m/%Y"),
                                      reparto.upper(), ruolo, note, filt)
                            self._save_history(out_file, nome, reparto.upper(), len(filt))
                            ok += 1
                            details.append(f"#{idx} OK · {out_file.name}")
                        except Exception as exc:  # noqa: BLE001
                            fail += 1
                            details.append(s_tr("bat_fail_generic", i=idx, e=str(exc)))
                    hashes = dict(saved_snap)
                    for t in tpl_snap:
                        try:
                            hashes[str(t.path)] = compute_template_hash(t.path)
                        except Exception:
                            pass
                    save_hashes(hashes)
                    self.saved_hashes = hashes
                finally:
                    self._post("done_batch", {
                        "ok": ok, "skip": skip, "fail": fail,
                        "detail": "\n".join(details[-30:]),
                        "out_dir": str(out_dir_snap),
                        "auto_open": bool(auto_open_batch),
                    })
                    self._post("worker_done", None)

            threading.Thread(target=work2, daemon=True).start()

        ttk.Button(top_row, text=self.tr("bat_choose"), style="Secondary.TButton",
                   command=load).pack(side=LEFT)

        footer = tk.Frame(body, bg=self._style_colors["card_body_bg"])
        footer.pack(fill=X, side="bottom")
        ttk.Label(footer, text=self.tr("bat_cols"), style="AppMuted.TLabel"
                  ).pack(side=LEFT)
        ttk.Button(footer, text=self.tr("bat_run"), style="Primary.TButton",
                   command=run).pack(side=RIGHT)

    def _parse_date_str(self, s: str):
        if not s:
            return None
        if isinstance(s, datetime):
            return s.date()
        if isinstance(s, date):
            return s
        s_norm = str(s).strip().replace("-", "/").replace(".", "/")
        for fmt in ("%d/%m/%Y", "%Y/%m/%d", "%d/%m/%y"):
            try:
                return datetime.strptime(s_norm, fmt).date()
            except ValueError:
                continue
        return None

    def _parse_batch_file(self, p: Path) -> list[dict[str, str]]:
        suffix = p.suffix.lower()
        if suffix in {".xlsx", ".xls"}:
            wb = load_workbook(p, data_only=True, read_only=True)
            try:
                ws = wb.worksheets[0]
                rows_iter = ws.iter_rows(values_only=True)
                raw_rows = [list(r) for r in rows_iter]
            finally:
                wb.close()
            if not raw_rows:
                return []
            header = [str(c).strip().lower() if c is not None else "" for c in raw_rows[0]]
            norm = []
            for c in header:
                if c in {"nome", "name"}:
                    norm.append("Nome")
                elif c in {"data", "date"}:
                    norm.append("Data")
                elif c in {"reparto", "department", "dipartimento"}:
                    norm.append("Reparto")
                elif c in {"ruolo", "role", "mansione"}:
                    norm.append("Ruolo")
                elif c in {"note", "notes"}:
                    norm.append("Note")
                else:
                    norm.append(c)
            out: list[dict[str, str]] = []
            for r in raw_rows[1:]:
                if all(v is None or str(v).strip() == "" for v in r):
                    continue
                obj: dict[str, str] = {}
                for key, val in zip(norm, r):
                    obj[key] = "" if val is None else str(val)
                out.append(obj)
            return out
        # CSV
        text = p.read_text(encoding="utf-8-sig", errors="replace")
        reader = csv.reader(text.splitlines())
        rows_list = list(reader)
        if not rows_list:
            return []
        header = [c.strip().lower() for c in rows_list[0]]
        norm = []
        for c in header:
            if c in {"nome", "name"}:
                norm.append("Nome")
            elif c in {"data", "date"}:
                norm.append("Data")
            elif c in {"reparto", "department", "dipartimento"}:
                norm.append("Reparto")
            elif c in {"ruolo", "role", "mansione"}:
                norm.append("Ruolo")
            elif c in {"note", "notes"}:
                norm.append("Note")
            else:
                norm.append(c)
        out = []
        for r in rows_list[1:]:
            if all(v.strip() == "" for v in r):
                continue
            while len(r) < len(norm):
                r.append("")
            out.append({k: v for k, v in zip(norm, r)})
        return out

    # ---------------------- History ---------------------------------------
    def _save_history(self, output_path: Path, name: str, department: str, total: int) -> None:
        history: list[dict[str, object]] = []
        if HISTORY_FILE.exists():
            try:
                history = json.loads(HISTORY_FILE.read_text(encoding="utf-8"))
            except (OSError, json.JSONDecodeError):
                history = []
        history.insert(0, {
            "path": str(output_path),
            "name": name,
            "department": department,
            "documents": total,
            "ts": datetime.now().isoformat(timespec="seconds"),
        })
        HISTORY_FILE.write_text(json.dumps(history[:50], ensure_ascii=False, indent=2),
                                encoding="utf-8")


# --------------------------- MAIN -----------------------------------------

def _enable_dpi_awareness() -> None:
    if os.name != "nt":
        return
    try:
        ctypes.windll.shcore.SetProcessDpiAwareness(1)  # PROCESS_PER_MONITOR_DPI_AWARE
    except Exception:
        try:
            ctypes.windll.user32.SetProcessDPIAware()
        except Exception:
            pass


def main() -> None:
    if tk is None:
        raise SystemExit(
            "Tkinter non è disponibile in Python. Installa una versione di Python "
            "con il supporto Tk per avviare l'interfaccia desktop."
        )
    _enable_dpi_awareness()
    root = tk.Tk()
    FormazioniApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
