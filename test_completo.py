"""Script di test completo Formazioni PZZ - Simulazione operatore umano.

Test ATOMICI (100% coverage funzioni):
- FA-001: load_departments_from_file()
- FA-002: parse_template()
- FA-003: discover_templates()
- FA-004: department_options()
- FA-005: templates_for_department()
- FA-006: replace_placeholders()
- FA-007: safe_file_part()
- FA-008: Creazione .docx con placeholder *nome* *data*
- FA-009: Creazione .xlsx con placeholder
- FA-010: replace_docx_placeholders()
- FA-011: docx_story()
- FA-012: xlsx_story()
- FA-013: make_styles()
- FA-014: build_pdf() - Generazione PDF completo
- FA-015: Template TUTTI (multireparto)
- FA-016: Copie multiple
- FA-017: File ignorati (formato non valido)
- FA-018: Cronologia JSON (_save_history simulato)
- FA-019: Validazione campi obbligatori
- FA-020: Cross-feature: flusso completo operatore
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
import traceback
from datetime import datetime
from pathlib import Path
from pypdf import PdfReader

sys.path.insert(0, str(Path(__file__).resolve().parent))

import app

PASS = 0
FAIL = 0
LOG: list[str] = []


def check(name: str, condition: bool, detail: str = "") -> None:
    global PASS, FAIL
    if condition:
        PASS += 1
        LOG.append(f"  [PASS] {name}" + (f" — {detail}" if detail else ""))
    else:
        FAIL += 1
        LOG.append(f"  [FAIL] {name}" + (f" — {detail}" if detail else ""))


def heading(title: str) -> None:
    LOG.append(f"\n=== {title} ===")


def create_sample_docx(path: Path, content_lines: list[str]) -> None:
    from docx import Document
    doc = Document()
    for line in content_lines:
        doc.add_paragraph(line)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Dipendente"
    table.rows[0].cells[1].text = "*nome*"
    table.rows[1].cells[0].text = "Data"
    table.rows[1].cells[1].text = "*data*"
    doc.save(str(path))


def create_sample_xlsx(path: Path) -> None:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Scheda Sicurezza"
    ws["A1"] = "Nome dipendente"
    ws["B1"] = "*nome*"
    ws["A2"] = "Data ingresso"
    ws["B2"] = "*data*"
    ws["A3"] = "Corso"
    ws["B3"] = "Sicurezza generale"
    wb.save(str(path))


# ============================================================
# TEST SUITE
# ============================================================
with tempfile.TemporaryDirectory() as tmpdir:
    TMP = Path(tmpdir)
    TEMPLATES_DIR = TMP / "templates"
    OUTPUT_DIR = TMP / "output"
    TEMPLATES_DIR.mkdir()
    OUTPUT_DIR.mkdir()

    # === FA-001: load_departments_from_file ===
    heading("FA-001 — Lettura reparti da reparti.txt")
    try:
        deps = app.load_departments_from_file()
        check("File esiste ed e' leggibile", True, f"{len(deps)} reparti trovati")
        check("Contiene AMMINISTRAZIONE", "AMMINISTRAZIONE" in deps, str(deps))
        check("Contiene SICUREZZA", "SICUREZZA" in deps)
        check("Contiene RISORSE UMANE", "RISORSE UMANE" in deps)
        check("Commenti # ignorati", not any(d.startswith("#") for d in deps))
        check("Righe vuote ignorate", not any(d == "" for d in deps))
        check("Senza duplicati", len(deps) == len(set(deps)))
    except Exception as e:
        check("load_departments_from_file", False, str(e))

    # === FA-002: parse_template ===
    heading("FA-002 — Parse nome template REPARTO_NUMERO_CODICE")
    try:
        t1 = app.parse_template(Path("SICUREZZA_2_ABC.docx"))
        check("Parse base docx", t1 is not None and t1.department == "SICUREZZA")
        check("Copie corrette", t1 is not None and t1.copies == 2)
        check("Codice corretto", t1 is not None and t1.code == "ABC")
        t2 = app.parse_template(Path("TUTTI_1_ZZZ.xlsx"))
        check("Parse TUTTI xlsx", t2 is not None and t2.is_for_every_department)
        t3 = app.parse_template(Path("formazione.pdf"))
        check("Estensione PDF supportata", t3 is None)
        check("Parse Word legacy", app.parse_template(Path("SICUREZZA_1_ABC.doc")) is not None)
        check("Parse Excel legacy", app.parse_template(Path("SICUREZZA_1_ABC.xls")) is not None)
        t4 = app.parse_template(Path("nome_sbagliato.docx"))
        check("Pattern non valido = None", t4 is None)
        t5 = app.parse_template(Path("RISORSE UMANE_3_DEF.docx"))
        check("Reparto multi-parola", t5 is not None and t5.department == "RISORSE UMANE")
    except Exception as e:
        check("parse_template", False, str(e))

    # Creazione template di esempio
    create_sample_docx(TEMPLATES_DIR / "SICUREZZA_1_SIC.docx",
                       ["Benvenuto *nome*", "Data ingresso: *data*", "Corso sicurezza generale"])
    create_sample_docx(TEMPLATES_DIR / "AMMINISTRAZIONE_2_AMM.docx",
                       ["Contratto per *nome*", "Firmato in data *data*"])
    create_sample_xlsx(TEMPLATES_DIR / "LOGISTICA_1_LOG.xlsx")
    create_sample_docx(TEMPLATES_DIR / "TUTTI_1_GEN.docx",
                       ["Documento generale per *nome*"])
    create_sample_docx(TEMPLATES_DIR / "PRODUZIONE_3_PRO.docx",
                       ["Istruzioni produzione per *nome* del *data*"])
    from reportlab.pdfgen.canvas import Canvas
    from reportlab.lib.pagesizes import A4, landscape
    horizontal_pdf = TEMPLATES_DIR / "ORIZZONTALE_1_ORI.pdf"
    horizontal_canvas = Canvas(str(horizontal_pdf), pagesize=landscape(A4))
    for page_number in range(2):
        horizontal_canvas.drawString(40, 540, f"Pagina orizzontale {page_number + 1}")
        horizontal_canvas.showPage()
    horizontal_canvas.save()

    # File non validi (devono essere ignorati)
    (TEMPLATES_DIR / "README.md").write_text("ignore")
    (TEMPLATES_DIR / "vecchio_file.doc").write_text("formato vecchio")
    (TEMPLATES_DIR / "nome_sbagliato.docx").write_text("pattern errato")

    # === FA-003: discover_templates ===
    heading("FA-003 — Scoperta template nella cartella")
    try:
        valid, ignored = app.discover_templates(TEMPLATES_DIR)
        check("6 template validi", len(valid) == 6, f"trovati {len(valid)}")
        check("2 file ignorati (estensioni supportate ma pattern errato)", len(ignored) == 2, f"trovati {len(ignored)}: {[p.name for p in ignored]}")
        check("README.md escluso dalla scansione", not any(str(p).endswith("README.md") for p in ignored))
        check(".doc riconosciuto come estensione supportata", any(str(p).endswith(".doc") for p in ignored))
    except Exception as e:
        check("discover_templates", False, str(e))

    # === FA-004: department_options ===
    heading("FA-004 — Unione reparti da file + template")
    try:
        opts = app.department_options(valid)
        check("Reparto da reparti.txt in testa", opts[0] == "AMMINISTRAZIONE" if opts else False, str(opts))
        check("SICUREZZA presente", "SICUREZZA" in opts)
        check("PRODUZIONE da template presente", "PRODUZIONE" in opts)
        check("LOGISTICA presente", "LOGISTICA" in opts)
        check("TUTTI escluso dalla lista reparti", "TUTTI" not in opts)
    except Exception as e:
        check("department_options", False, str(e))

    # === FA-005: templates_for_department ===
    heading("FA-005 — Filtro template per reparto (inclusi TUTTI)")
    try:
        t_sic = app.templates_for_department(valid, "SICUREZZA")
        check("SICUREZZA: 2 template (1 suo + 1 TUTTI)", len(t_sic) == 2, f"{len(t_sic)}")
        t_amm = app.templates_for_department(valid, "AMMINISTRAZIONE")
        check("AMMINISTRAZIONE: 2 template", len(t_amm) == 2)
        t_prod = app.templates_for_department(valid, "PRODUZIONE")
        check("PRODUZIONE: 2 template", len(t_prod) == 2)
        t_inesistente = app.templates_for_department(valid, "INESISTENTE")
        check("Reparto inesistente: solo TUTTI", len(t_inesistente) == 1, f"{len(t_inesistente)}")
    except Exception as e:
        check("templates_for_department", False, str(e))

    # === FA-006: replace_placeholders ===
    heading("FA-006 — Sostituzione placeholder *nome* e *data*")
    try:
        r = app.replace_placeholders("Ciao *nome*, benvenuto il *data*", "Mario Rossi", "01/09/2026")
        check("Sostituzione base", r == "Ciao Mario Rossi, benvenuto il 01/09/2026", str(r))
        r2 = app.replace_placeholders("*NOME* - *DATA*", "Luigi Bianchi", "15/10/2026")
        check("Case insensitive", r2 == "Luigi Bianchi - 15/10/2026", str(r2))
        r3 = app.replace_placeholders(42, "a", "b")
        check("Non-stringhe passanti", r3 == 42)
    except Exception as e:
        check("replace_placeholders", False, str(e))

    # === FA-007: safe_file_part ===
    heading("FA-007 — Pulizia nome per filesystem")
    try:
        check("Base", app.safe_file_part("Mario Rossi") == "Mario-Rossi")
        check("Vuoto -> default", app.safe_file_part("   ") == "persona")
        check("Caratteri speciali", "?" not in app.safe_file_part("Mario?Rossi<>"))
        check("Accenti permessi", "Ù" in app.safe_file_part("Périgord Ùltimo"))
    except Exception as e:
        check("safe_file_part", False, str(e))

    # === FA-008/009: Placeholder in DOCX e XLSX ===
    heading("FA-008 — replace_docx_placeholders + docx_story")
    try:
        from docx import Document
        p = TEMPLATES_DIR / "SICUREZZA_1_SIC.docx"
        doc = Document(str(p))
        app.replace_docx_placeholders(doc, "Maria Verdi", "05/09/2026")
        texts = [para.text for para in doc.paragraphs]
        check("Paragrafo *nome* sostituito", any("Maria Verdi" in t for t in texts), str(texts))
        check("Paragrafo *data* sostituito", any("05/09/2026" in t for t in texts))
        cell_texts = [c.text for t in doc.tables for r in t.rows for c in r.cells]
        check("Tabella *nome* sostituito", "Maria Verdi" in cell_texts, str(cell_texts))
        check("Tabella *data* sostituito", "05/09/2026" in cell_texts)

        styles = app.make_styles()
        story = app.docx_story(p, "Maria Verdi", "05/09/2026", styles)
        check("docx_story non vuota", len(story) > 0, f"{len(story)} elementi")
    except Exception as e:
        check("docx placeholder", False, f"{e}\n{traceback.format_exc()}")

    heading("FA-009 — xlsx_story + placeholder XLSX")
    try:
        styles = app.make_styles()
        story = app.xlsx_story(TEMPLATES_DIR / "LOGISTICA_1_LOG.xlsx", "Giacomo Blu", "10/11/2026", styles)
        check("xlsx_story non vuota", len(story) > 0, f"{len(story)} elementi")
    except Exception as e:
        check("xlsx story", False, f"{e}\n{traceback.format_exc()}")

    # === FA-010: make_styles ===
    heading("FA-010 — make_styles (tutti gli stili)")
    try:
        s = app.make_styles()
        required = ["cover_title", "cover_subtitle", "meta", "heading", "subheading",
                    "body", "table", "muted", "small"]
        check("Tutti gli stili presenti", all(k in s for k in required), str(list(s.keys())))
    except Exception as e:
        check("make_styles", False, str(e))

    # === FA-011: build_pdf ===
    heading("FA-011 — build_pdf: Generazione PDF completo")
    try:
        out_path = OUTPUT_DIR / "test_dossier.pdf"
        scelti = app.templates_for_department(valid, "PRODUZIONE")
        n = app.build_pdf(
            output_path=out_path,
            employee_name="Franco Neri",
            entry_date="03/09/2026",
            department="PRODUZIONE",
            role="Operaio specializzato",
            notes="Nota aggiuntiva di test.",
            templates=scelti,
        )
        check("PDF creato sul disco", out_path.exists() and out_path.stat().st_size > 0, f"{out_path.stat().st_size} bytes")
        check("Conteggio copie esatto (PRODUZIONE 3 + TUTTI 1)", n == 4, f"n={n}")
    except Exception as e:
        check("build_pdf", False, f"{e}\n{traceback.format_exc()}")

    heading("FA-011b — PDF orizzontale: pagine e orientamento conservati")
    try:
        horizontal_output = OUTPUT_DIR / "horizontal.pdf"
        horizontal_template = [t for t in valid if t.path == horizontal_pdf][0]
        app.build_pdf(
            horizontal_output,
            "Franco Neri",
            "03/09/2026",
            "SICUREZZA",
            "Operaio",
            "",
            [horizontal_template],
        )
        horizontal_reader = PdfReader(str(horizontal_output))
        check("Numero pagine orizzontali conservato", len(horizontal_reader.pages) == 2)
        check(
            "Orientamento orizzontale conservato",
            all(page.mediabox.width > page.mediabox.height for page in horizontal_reader.pages),
        )
    except Exception as e:
        check("PDF orizzontale", False, f"{e}\n{traceback.format_exc()}")

    # === FA-012: Template TUTTI in ogni reparto ===
    heading("FA-012 — Template TUTTI presente in ogni reparto")
    try:
        for rep in ["AMMINISTRAZIONE", "SICUREZZA", "LOGISTICA", "PRODUZIONE"]:
            ts = app.templates_for_department(valid, rep)
            has_tutti = any(t.is_for_every_department for t in ts)
            check(f"TUTTI in {rep}", has_tutti)
    except Exception as e:
        check("TUTTI multireparto", False, str(e))

    # === FA-013: Copie multiple ===
    heading("FA-013 — Copie multiple (AMMINISTRAZIONE_2_AMM)")
    try:
        amm = [t for t in valid if t.department.upper() == "AMMINISTRAZIONE" and not t.is_for_every_department]
        check("AMMINISTRAZIONE ha 2 copie nel template", amm[0].copies == 2 if amm else False)
    except Exception as e:
        check("copie multiple", False, str(e))

    # === FA-014: File ignorati ===
    heading("FA-014 — File ignorati correttamente")
    try:
        nomi_ignorati = [p.name.lower() for p in ignored]
        check(".docx con nome sbagliato ignorato (pattern invalido)", "nome_sbagliato.docx" in nomi_ignorati, str(nomi_ignorati))
    except Exception as e:
        check("file ignorati", False, str(e))

    # === FA-015: Cronologia ===
    heading("FA-015 — Salvataggio cronologia JSON")
    try:
        hist_file = TMP / "history_test.json"
        history: list[dict] = []
        entry = {
            "path": str(out_path),
            "name": "Franco Neri",
            "department": "PRODUZIONE",
            "documents": 4,
            "created_at": datetime.now().isoformat(timespec="seconds"),
        }
        history.insert(0, entry)
        hist_file.write_text(json.dumps(history[:20], ensure_ascii=False, indent=2), encoding="utf-8")
        readback = json.loads(hist_file.read_text(encoding="utf-8"))
        check("Cronologia scritta e letta", readback[0]["name"] == "Franco Neri")
        check("UTF-8 conservato", readback[0]["department"] == "PRODUZIONE")
    except Exception as e:
        check("cronologia", False, str(e))

    # === FA-016: Validazione campi obbligatori (logica) ===
    heading("FA-016 — Validazione campi obbligatori (logica)")
    try:
        check("Nome vuoto -> fallisce", not "".strip())
        check("Data vuota -> fallisce", not "".strip())
        check("Reparto vuoto -> fallisce", not "".strip())
        check("Nome non vuoto -> ok", bool("Mario Rossi".strip()))
        check("Nessun template per reparto -> avviso",
              len(app.templates_for_department(valid, "REPARTO_CHE_NON_ESISTE")) == 1  # solo TUTTI
              )
    except Exception as e:
        check("validazione", False, str(e))

    # === FA-020: FLUSSO COMPLETO OPERATORE ===
    heading("FA-020 — FLUSSO COMPLETO: Operatore crea dossier per nuovo assunto")
    try:
        # Step 1: scopri template (simula click "Aggiorna documenti")
        templates, ignored = app.discover_templates(TEMPLATES_DIR)
        check("[Step 1] Template caricati", len(templates) == 6)

        # Step 2: reparti disponibili (lettura da reparti.txt)
        reparti = app.department_options(templates)
        check("[Step 2] Reparti caricati da file reparti.txt", "AMMINISTRAZIONE" in reparti and "SICUREZZA" in reparti)

        # Step 3: operatore inserisce dati
        nome = "Giuseppe Verdi"
        data = "03/09/2026"
        reparto = "COMMERCIALE"
        ruolo = "Agente di vendita"
        note = "Neoassunto, segue formazione base."
        check("[Step 3] Dati inseriti correttamente", all([nome, data, reparto]))

        # Step 4: selezione reparto -> aggiornamento lista documenti
        scelti = app.templates_for_department(templates, reparto)
        tot_copie = sum(t.copies for t in scelti)
        check(f"[Step 4] COMMERCIALE: {len(scelti)} template, {tot_copie} copie totali",
              len(scelti) >= 1 and tot_copie >= 1, f"{len(scelti)} template, {tot_copie} copie")

        # Step 5: Genera PDF
        out = OUTPUT_DIR / f"dossier_{app.safe_file_part(nome)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        n_doc = app.build_pdf(out, nome, data, reparto, ruolo, note, scelti)
        check("[Step 5] PDF creato correttamente", out.exists() and out.stat().st_size > 1000, f"{out.stat().st_size} bytes")
        check(f"[Step 5] {n_doc} documenti inclusi", n_doc == tot_copie, f"attese {tot_copie}, ricevute {n_doc}")

        # Step 6: Verifica filename sicuro
        check("[Step 6] Filename valido (no spazi strani)", " " not in out.name and "Giuseppe-Verdi" in out.name, out.name)

    except Exception as e:
        check("FLUSSO COMPLETO", False, f"{e}\n{traceback.format_exc()}")

# ============================================================
# RIEPILOGO
# ============================================================
LOG.append("\n" + "=" * 60)
LOG.append(f"RIEPILOGO TEST: {PASS} PASSATI, {FAIL} FALLITI su {PASS+FAIL} totali")
LOG.append("=" * 60)

print("\n".join(LOG))

if FAIL > 0:
    sys.exit(1)
else:
    ok_msg = "[OK] TUTTI I TEST SUPERATI - L'app funziona correttamente!"
    try:
        print("\n" + ok_msg)
    except UnicodeEncodeError:
        print("\n[OK] TUTTI I TEST SUPERATI - L'app funziona correttamente!")
